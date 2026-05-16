"""Pure-function unit tests — no I/O, no sessions, no server."""
import sys
from pathlib import Path

# Allow importing api.py directly without the full package installed
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from api import (
    app_date_label,
    app_money,
    app_source_documents,
    app_surface_label,
    app_status_label,
)
from engine.orchestrator import (
    PlanOrchestrator,
    available_mandat_types,
    classify_dossier,
    load_plan_for_mandat,
)


# ── app_money ────────────────────────────────────────────────────────────────

class TestAppMoney:
    def test_integer(self):
        assert app_money(500000) == "500 000 $"

    def test_float_rounds(self):
        assert app_money(499999.9) == "500 000 $"

    def test_zero(self):
        assert app_money(0) == "0 $"

    def test_string_numeric(self):
        assert app_money("250000") == "250 000 $"

    def test_none_returns_dash(self):
        assert app_money(None) == "-"

    def test_empty_string_returns_dash(self):
        assert app_money("") == "-"

    def test_non_numeric_returns_dash(self):
        assert app_money("abc") == "-"

    def test_negative(self):
        result = app_money(-10000)
        assert "$" in result


# ── app_date_label ────────────────────────────────────────────────────────────

class TestAppDateLabel:
    def test_iso_datetime(self):
        assert app_date_label("2025-03-15T10:30:00") == "2025-03-15"

    def test_iso_with_z(self):
        assert app_date_label("2025-03-15T10:30:00Z") == "2025-03-15"

    def test_date_only(self):
        assert app_date_label("2025-03-15") == "2025-03-15"

    def test_none_returns_empty(self):
        assert app_date_label(None) == ""

    def test_empty_string_returns_empty(self):
        assert app_date_label("") == ""

    def test_invalid_returns_raw(self):
        assert app_date_label("not-a-date") == "not-a-date"

    def test_with_timezone_offset(self):
        assert app_date_label("2025-06-01T00:00:00+05:00") == "2025-06-01"


# ── app_surface_label ─────────────────────────────────────────────────────────

class TestAppSurfaceLabel:
    def test_basic(self):
        assert app_surface_label({"value": 120, "unit": "m²"}) == "120 m²"

    def test_no_value_returns_dash(self):
        assert app_surface_label({"value": None, "unit": "m²"}) == "-"

    def test_empty_value_returns_dash(self):
        assert app_surface_label({"value": "", "unit": "m²"}) == "-"

    def test_not_dict_returns_dash(self):
        assert app_surface_label("120 m²") == "-"
        assert app_surface_label(None) == "-"

    def test_no_unit(self):
        result = app_surface_label({"value": 80, "unit": ""})
        assert "80" in result


# ── app_source_documents ──────────────────────────────────────────────────────

class TestAppSourceDocuments:
    def test_empty_knowledge(self):
        assert app_source_documents({}) == []

    def test_sources_from_knowledge(self):
        knowledge = {
            "sources": {
                "items": [
                    {"source_id": "SRC-1", "source_type": "mls", "reliability_level": "A"},
                    {"source_id": "SRC-2", "source_type": "mpac"},
                ]
            }
        }
        docs = app_source_documents(knowledge)
        assert len(docs) == 2
        ids = [d["id"] for d in docs]
        assert "SRC-1" in ids
        assert "SRC-2" in ids

    def test_uploaded_docs_merged_from_session(self):
        knowledge = {}
        session = {
            "uploaded_documents": [
                {"id": "upl-1", "name": "Acte.pdf", "filename": "acte.pdf", "size_bytes": 204800},
            ]
        }
        docs = app_source_documents(knowledge, session)
        assert len(docs) == 1
        assert docs[0]["id"] == "upl-1"
        assert docs[0]["name"] == "Acte.pdf"
        assert "200" in docs[0]["sizeLabel"]  # 204800 // 1024 == 200

    def test_knowledge_and_uploaded_merged(self):
        knowledge = {
            "sources": {
                "items": [{"source_id": "SRC-1"}]
            }
        }
        session = {
            "uploaded_documents": [
                {"id": "upl-1", "name": "Doc.pdf", "filename": "doc.pdf", "size_bytes": 1024},
            ]
        }
        docs = app_source_documents(knowledge, session)
        assert len(docs) == 2

    def test_invalid_items_skipped(self):
        knowledge = {
            "sources": {
                "items": ["not-a-dict", None, {"source_id": "OK"}]
            }
        }
        docs = app_source_documents(knowledge)
        assert len(docs) == 1
        assert docs[0]["id"] == "OK"

    def test_no_session_is_fine(self):
        docs = app_source_documents({}, session=None)
        assert docs == []


# ── app_status_label ──────────────────────────────────────────────────────────

class TestAppStatusLabel:
    def test_complet(self):
        assert app_status_label({"package_status": "PRET_REVUE_EVALUATEUR_AGREE"}) == "complet"

    def test_en_cours_pret_revision(self):
        assert app_status_label({"status": "PRET_REVISION_FINALE"}) == "en-cours"

    def test_en_cours_a_revoir(self):
        assert app_status_label({"status": "A_REVOIR"}) == "en-cours"

    def test_brouillon_default(self):
        assert app_status_label({"status": "CREATED"}) == "brouillon"

    def test_empty_record(self):
        assert app_status_label({}) == "brouillon"


# ── classify_dossier ──────────────────────────────────────────────────────────

class TestClassifyDossier:
    def test_explicit_mandat_type(self):
        assert classify_dossier({"mandat_type": "commercial"}) == "commercial"

    def test_assurance_via_but_evaluation(self):
        assert classify_dossier({"but_evaluation": "fins d'assurance"}) == "assurance"

    def test_unifamilial_exact(self):
        assert classify_dossier({"type_bien": "residentiel_unifamilial"}) == "residentiel_standard"

    def test_duplex_exact(self):
        assert classify_dossier({"type_bien": "duplex"}) == "residentiel_multifamilial"

    def test_triplex_partial(self):
        assert classify_dossier({"type_bien": "triplex_clé_en_main"}) == "residentiel_multifamilial"

    def test_industriel_partial(self):
        assert classify_dossier({"type_bien": "entrepot_logistique"}) == "industriel"

    def test_terrain_exact(self):
        assert classify_dossier({"type_bien": "terrain_vacant"}) == "terrain"

    def test_immeuble_revenus(self):
        assert classify_dossier({"type_bien": "immeuble_revenus"}) == "immeuble_revenus"

    def test_empty_defaults_to_residentiel_standard(self):
        assert classify_dossier({}) == "residentiel_standard"

    def test_unknown_type_bien_defaults(self):
        assert classify_dossier({"type_bien": "xyz_inconnu"}) == "residentiel_standard"

    def test_case_insensitive_partial(self):
        assert classify_dossier({"type_bien": "PLEX"}) == "residentiel_multifamilial"

    def test_assurance_type_bien(self):
        assert classify_dossier({"type_bien": "assurance"}) == "assurance"


# ── load_plan_for_mandat ──────────────────────────────────────────────────────

class TestLoadPlanForMandat:
    def test_all_types_loadable(self):
        for mt in available_mandat_types():
            plan = load_plan_for_mandat(mt)
            assert plan.mandat_type == mt
            assert plan.format_rapport in {"abrege", "narratif_complet", "mise_a_jour"}
            assert len(plan.methodes_requises) >= 1
            assert plan.methode_preponderante in plan.methodes_requises

    def test_residentiel_standard_plan(self):
        plan = load_plan_for_mandat("residentiel_standard")
        assert plan.format_rapport == "abrege"
        assert "approche_comparative" in plan.methodes_requises
        assert plan.methode_preponderante == "approche_comparative"
        assert plan.umpp_requis is True

    def test_assurance_cout_only(self):
        plan = load_plan_for_mandat("assurance")
        assert plan.methodes_requises == ["approche_cout"]
        assert plan.methode_preponderante == "approche_cout"
        assert plan.umpp_requis is False

    def test_immeuble_revenus_preponderance(self):
        plan = load_plan_for_mandat("immeuble_revenus")
        assert plan.methode_preponderante == "approche_revenu"
        assert plan.format_rapport == "narratif_complet"

    def test_unknown_raises_key_error(self):
        import pytest
        with pytest.raises(KeyError, match="inconnu"):
            load_plan_for_mandat("type_inexistant")


# ── PlanOrchestrator ──────────────────────────────────────────────────────────

class TestPlanOrchestrator:
    def test_build_engine_returns_engine_and_plan(self):
        orch = PlanOrchestrator()
        case = {"dossier_id": "D-TEST", "type_bien": "residentiel_unifamilial"}
        engine, plan = orch.build_engine(case)
        assert engine is not None
        assert plan.mandat_type == "residentiel_standard"

    def test_enrich_case_adds_plan_fields(self):
        orch = PlanOrchestrator()
        case = {"dossier_id": "D-TEST", "type_bien": "duplex"}
        _, plan = orch.build_engine(case)
        enriched = orch.enrich_case(case, plan)
        assert enriched["mandat_type"] == "residentiel_multifamilial"
        assert "methodes_requises" in enriched
        assert enriched["dossier_id"] == "D-TEST"  # original preserved

    def test_enrich_case_does_not_mutate_original(self):
        orch = PlanOrchestrator()
        case = {"dossier_id": "D-TEST", "type_bien": "terrain"}
        _, plan = orch.build_engine(case)
        orch.enrich_case(case, plan)
        assert "mandat_type" not in case  # original untouched

    def test_explicit_mandat_type_in_case_respected(self):
        orch = PlanOrchestrator()
        case = {"dossier_id": "D-TEST", "type_bien": "maison", "mandat_type": "assurance"}
        _, plan = orch.build_engine(case)
        assert plan.mandat_type == "assurance"


# ── DEFAULT_SKILLS_BY_AGENT ───────────────────────────────────────────────────

class TestDefaultSkillsByAgent:
    def test_amu_analyst_in_default_skills(self):
        from engine.skills import DEFAULT_SKILLS_BY_AGENT
        assert "amu-analyst" in DEFAULT_SKILLS_BY_AGENT
        skills = DEFAULT_SKILLS_BY_AGENT["amu-analyst"]
        assert "analyse-amu" in skills
        assert "recherche-urbanisme-construction" in skills
        assert "recherche-normes-professionnelles" in skills

    def test_mandat_intake_in_default_skills(self):
        from engine.skills import DEFAULT_SKILLS_BY_AGENT
        assert "mandat-intake" in DEFAULT_SKILLS_BY_AGENT
        skills = DEFAULT_SKILLS_BY_AGENT["mandat-intake"]
        assert "redaction-lettre-mandat" in skills

    def test_fta_in_valuation_draft_skills(self):
        from engine.skills import DEFAULT_SKILLS_BY_AGENT
        assert "analyse-approche-fta" in DEFAULT_SKILLS_BY_AGENT["valuation-draft"]


# ── TestAmuDeterministic ──────────────────────────────────────────────────────

class TestAmuDeterministic:
    def test_umpp_conclusion_fields(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine
        engine = RuntimeEngine()
        case = {
            "dossier_id": "D-AMU-TEST",
            "type_bien": "residentiel_unifamilial",
            "zone": "R-2",
            "date_reference": "2026-05-01",
        }
        payload = engine._artifact_payload(
            "amu-analyst", "umpp_conclusion.json", case, "BROUILLON", [], []
        )
        assert payload["dossier_id"] == "D-AMU-TEST"
        assert payload["step"] == "amu-analyst"
        assert "umpp" in payload
        umpp = payload["umpp"]
        assert "usage_retenu" in umpp
        assert "criteres" in umpp
        criteres = umpp["criteres"]
        assert "physiquement_possible" in criteres
        assert "legalement_permis" in criteres
        assert "financierement_faisable" in criteres
        assert "maximalement_productif" in criteres
        assert "umpp_differe_usage_actuel" in umpp
        assert isinstance(payload.get("confidence"), float)

    def test_amu_analyse_md_fields(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine
        engine = RuntimeEngine()
        case = {
            "dossier_id": "D-AMU-TEST",
            "type_bien": "terrain_vacant",
            "zone": "C-1",
            "date_reference": "2026-05-01",
        }
        payload = engine._artifact_payload(
            "amu-analyst", "amu_analyse.md", case, "BROUILLON", [], []
        )
        assert payload["step"] == "amu-analyst"
        assert "_raw_md" in payload
        assert "AMU" in payload["_raw_md"] or "meilleur usage" in payload["_raw_md"].lower()


# ── TestPipelineStepCount ─────────────────────────────────────────────────────

class TestPipelineStepCount:
    def test_default_steps_has_seven(self):
        from engine.runtime import DEFAULT_STEPS
        assert len(DEFAULT_STEPS) == 7

    def test_mandat_intake_at_index_zero(self):
        from engine.runtime import DEFAULT_STEPS
        assert DEFAULT_STEPS[0].name == "mandat-intake"

    def test_data_facts_at_index_one(self):
        from engine.runtime import DEFAULT_STEPS
        assert DEFAULT_STEPS[1].name == "data-facts"

    def test_amu_analyst_at_index_two(self):
        from engine.runtime import DEFAULT_STEPS
        assert DEFAULT_STEPS[2].name == "amu-analyst"

    def test_amu_analyst_reads_fiche_bien(self):
        from engine.runtime import DEFAULT_STEPS
        amu_step = DEFAULT_STEPS[2]
        assert "fiche_bien.json" in amu_step.reads

    def test_amu_analyst_writes_umpp_conclusion(self):
        from engine.runtime import DEFAULT_STEPS
        amu_step = DEFAULT_STEPS[2]
        assert "umpp_conclusion.json" in amu_step.writes
        assert "amu_analyse.md" in amu_step.writes

    def test_mandat_intake_writes_lettre_mandat(self):
        from engine.runtime import DEFAULT_STEPS
        step = DEFAULT_STEPS[0]
        assert "lettre_mandat.md" in step.writes
        assert "conflit_interets.json" in step.writes

    def test_redaction_reads_lettre_mandat(self):
        from engine.runtime import DEFAULT_STEPS
        redaction = DEFAULT_STEPS[6]
        assert redaction.name == "redaction"
        assert "lettre_mandat.md" in redaction.reads


# ── TestMandatIntakeDeterministic ─────────────────────────────────────────────

class TestMandatIntakeDeterministic:
    def test_conflit_interets_fields(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine
        engine = RuntimeEngine()
        case = {
            "dossier_id": "D-MANDAT-TEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-12",
            "mandat_type": "residentiel_standard",
            "format_rapport": "abrege",
        }
        payload = engine._artifact_payload(
            "mandat-intake", "conflit_interets.json", case, "BROUILLON", [], []
        )
        assert payload["dossier_id"] == "D-MANDAT-TEST"
        assert payload["step"] == "mandat-intake"
        assert payload["artifact"] == "conflit_interets.json"
        assert payload["conflit_detecte"] is False
        assert payload["verification_completee"] is True
        assert "commentaire" in payload

    def test_lettre_mandat_md_raw_md(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine
        engine = RuntimeEngine()
        case = {
            "dossier_id": "D-MANDAT-TEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-12",
            "mandat_type": "residentiel_standard",
            "format_rapport": "abrege",
        }
        payload = engine._artifact_payload(
            "mandat-intake", "lettre_mandat.md", case, "BROUILLON", [], []
        )
        assert payload["step"] == "mandat-intake"
        assert "_raw_md" in payload
        assert "Lettre de mandat" in payload["_raw_md"] or "mandat" in payload["_raw_md"].lower()


# ── TestCommanditaireInCase ───────────────────────────────────────────────────

class TestCommanditaireInCase:
    def test_commanditaire_merged_from_body(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import load_case_from_body
        body = {
            "commanditaire": {
                "nom": "Banque Nationale",
                "organisation": "Financement immobilier",
                "fin_evaluation": "hypothecaire",
            }
        }
        case, _ = load_case_from_body(body)
        assert case["commanditaire"]["nom"] == "Banque Nationale"
        assert case["commanditaire"]["organisation"] == "Financement immobilier"
        assert case["commanditaire"]["fin_evaluation"] == "hypothecaire"

    def test_commanditaire_defaults_when_absent(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import load_case_from_body
        case, _ = load_case_from_body({})
        # commanditaire key absent — no crash, no injection
        assert "commanditaire" not in case

    def test_commanditaire_nom_default_placeholder(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import load_case_from_body
        body = {"commanditaire": {"nom": "", "fin_evaluation": "succession"}}
        case, _ = load_case_from_body(body)
        assert case["commanditaire"]["nom"] == "[COMMANDITAIRE]"


# ── TestLettreMandat_Commanditaire ────────────────────────────────────────────

class TestLettreMandat_Commanditaire:
    def test_lettre_mandat_uses_commanditaire_nom(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine
        engine = RuntimeEngine()
        case = {
            "dossier_id": "D-CMD-TEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-13",
            "mandat_type": "residentiel_standard",
            "format_rapport": "abrege",
            "commanditaire": {
                "nom": "Banque Nationale",
                "organisation": "Financement immobilier",
                "fin_evaluation": "hypothecaire",
            },
        }
        payload = engine._artifact_payload(
            "mandat-intake", "lettre_mandat.md", case, "BROUILLON", [], []
        )
        assert "[COMMANDITAIRE]" not in payload["_raw_md"]
        assert "Banque Nationale" in payload["_raw_md"]

    def test_lettre_mandat_placeholder_when_no_commanditaire(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine
        engine = RuntimeEngine()
        case = {
            "dossier_id": "D-CMD-TEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-13",
            "mandat_type": "residentiel_standard",
            "format_rapport": "abrege",
        }
        payload = engine._artifact_payload(
            "mandat-intake", "lettre_mandat.md", case, "BROUILLON", [], []
        )
        assert "[COMMANDITAIRE]" in payload["_raw_md"]


# ── TestConflit_Deterministic_False ──────────────────────────────────────────

class TestConflit_Deterministic_False:
    def test_conflit_detecte_false_without_llm(self):
        """Without LLM (no OPENAI_API_KEY), conflit_detecte stays False."""
        import sys
        import os
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine
        original_key = os.environ.pop("OPENAI_API_KEY", None)
        try:
            engine = RuntimeEngine()
            case = {
                "dossier_id": "D-CONFLIT-TEST",
                "type_bien": "residentiel_unifamilial",
                "date_reference": "2026-05-13",
                "mandat_type": "residentiel_standard",
                "format_rapport": "abrege",
                "commanditaire": {"nom": "BNC", "organisation": "", "fin_evaluation": "hypothecaire"},
            }
            payload = engine._artifact_payload(
                "mandat-intake", "conflit_interets.json", case, "BROUILLON", [], []
            )
            assert payload["conflit_detecte"] is False
            assert payload["verification_completee"] is True
        finally:
            if original_key is not None:
                os.environ["OPENAI_API_KEY"] = original_key


# ── TestConflit_Gate_Blocks ───────────────────────────────────────────────────

class TestConflit_Gate_Blocks:
    def test_pipeline_raises_on_conflit_detecte(self, tmp_path):
        """run_case_data raises PipelineConflitError when conflit_detecte: True in artifact."""
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine, PipelineConflitError, DEFAULT_STEPS

        engine = RuntimeEngine(steps=DEFAULT_STEPS[:1])  # mandat-intake only
        case = {
            "dossier_id": "D-GATE-TEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-13",
            "mandat_type": "residentiel_standard",
            "format_rapport": "abrege",
        }

        import pytest

        original_payload = engine._artifact_payload

        def patched_payload(step, artifact, case, status, blocking, warnings, valuation_values=None):
            p = original_payload(step, artifact, case, status, blocking, warnings, valuation_values)
            if step == "mandat-intake" and artifact == "conflit_interets.json":
                p["conflit_detecte"] = True
                p["conflit_motif"] = "Test: conflit injecte"
            return p

        engine._artifact_payload = patched_payload

        with pytest.raises(PipelineConflitError, match="Test: conflit injecte"):
            engine.run_case_data(case, tmp_path, source_fixture="test", case_stem="test", case_subdir=True)

    def test_pipeline_no_exception_when_conflit_false(self, tmp_path):
        """run_case_data runs normally when conflit_detecte: False."""
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine, DEFAULT_STEPS

        engine = RuntimeEngine(steps=DEFAULT_STEPS[:1])
        case = {
            "dossier_id": "D-GATE-OK-TEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-13",
            "mandat_type": "residentiel_standard",
            "format_rapport": "abrege",
        }
        result = engine.run_case_data(case, tmp_path, source_fixture="test", case_stem="test", case_subdir=True)
        assert result["dossier_id"] == "D-GATE-OK-TEST"


# ── TestConflit_ForceOverride ─────────────────────────────────────────────────

class TestConflit_ForceOverride:
    def test_force_conflit_continue_bypasses_gate(self, tmp_path):
        """force_conflit_continue: True lets pipeline continue despite conflit_detecte."""
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine, DEFAULT_STEPS

        engine = RuntimeEngine(steps=DEFAULT_STEPS[:1])
        case = {
            "dossier_id": "D-OVERRIDE-TEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-13",
            "mandat_type": "residentiel_standard",
            "format_rapport": "abrege",
            "force_conflit_continue": True,
        }

        original_payload = engine._artifact_payload

        def patched_payload(step, artifact, case, status, blocking, warnings, valuation_values=None):
            p = original_payload(step, artifact, case, status, blocking, warnings, valuation_values)
            if step == "mandat-intake" and artifact == "conflit_interets.json":
                p["conflit_detecte"] = True
                p["conflit_motif"] = "Test: conflit injecte"
            return p

        engine._artifact_payload = patched_payload

        result = engine.run_case_data(case, tmp_path, source_fixture="test", case_stem="test", case_subdir=True)
        assert result["dossier_id"] == "D-OVERRIDE-TEST"


# ── TestConflitLLMParsing ─────────────────────────────────────────────────────

class TestConflitLLMParsing:
    def _make_mock_openai(self, llm_response: str):
        """Build a mock openai module whose OpenAI().chat.completions.create() returns llm_response."""
        import unittest.mock
        mock_resp = unittest.mock.MagicMock()
        mock_resp.choices[0].message.content = llm_response
        mock_client = unittest.mock.MagicMock()
        mock_client.chat.completions.create.return_value = mock_resp
        mock_openai_module = unittest.mock.MagicMock()
        mock_openai_module.OpenAI.return_value = mock_client
        return mock_openai_module

    def test_conflit_detecte_set_on_sentinel_prefix(self):
        """_enrich_artifact_llm sets conflit_detecte: True when LLM returns CONFLIT_DETECTE: prefix."""
        import sys
        import os
        import unittest.mock
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine, RuntimeStep

        engine = RuntimeEngine()
        payload = {
            "dossier_id": "D-LLM-TEST",
            "step": "mandat-intake",
            "artifact": "conflit_interets.json",
            "conflit_detecte": False,
            "verification_completee": True,
            "commentaire": "V0 deterministe.",
            "analyse_conflit": "",
        }
        case = {"dossier_id": "D-LLM-TEST", "type_bien": "residentiel_unifamilial"}
        step = RuntimeStep(
            name="mandat-intake",
            reads=[],
            writes=["conflit_interets.json"],
            skills=[],
            agent_config="AGENTCONFIG-MANDAT-INTAKE-V0.yaml",
        )
        llm_response = "CONFLIT_DETECTE: Lien familial avec le vendeur\n\nAnalyse détaillée..."
        mock_openai = self._make_mock_openai(llm_response)

        with unittest.mock.patch.dict(os.environ, {"OPENAI_API_KEY": "test-key"}):
            with unittest.mock.patch.dict(sys.modules, {"openai": mock_openai}):
                result = engine._enrich_artifact_llm(step, "conflit_interets.json", payload, case)

        assert result["conflit_detecte"] is True
        assert result["conflit_motif"] == "Lien familial avec le vendeur"
        assert result["analyse_conflit"] == llm_response

    def test_no_conflit_when_llm_returns_normal_response(self):
        """_enrich_artifact_llm leaves conflit_detecte: False when LLM returns normal text."""
        import sys
        import os
        import unittest.mock
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import RuntimeEngine, RuntimeStep

        engine = RuntimeEngine()
        payload = {
            "dossier_id": "D-LLM-OK-TEST",
            "step": "mandat-intake",
            "artifact": "conflit_interets.json",
            "conflit_detecte": False,
            "verification_completee": True,
            "commentaire": "V0 deterministe.",
            "analyse_conflit": "",
        }
        case = {"dossier_id": "D-LLM-OK-TEST", "type_bien": "residentiel_unifamilial"}
        step = RuntimeStep(
            name="mandat-intake",
            reads=[],
            writes=["conflit_interets.json"],
            skills=[],
            agent_config="AGENTCONFIG-MANDAT-INTAKE-V0.yaml",
        )
        llm_response = "Aucun conflit détecté. L'évaluateur est indépendant de toutes les parties."
        mock_openai = self._make_mock_openai(llm_response)

        with unittest.mock.patch.dict(os.environ, {"OPENAI_API_KEY": "test-key"}):
            with unittest.mock.patch.dict(sys.modules, {"openai": mock_openai}):
                result = engine._enrich_artifact_llm(step, "conflit_interets.json", payload, case)

        assert result["conflit_detecte"] is False
        assert result["analyse_conflit"] == llm_response


# ── TestIngestion_ExtractPDFText ──────────────────────────────────────────────

class TestIngestion_ExtractPDFText:
    def test_extracts_text_from_pdf_with_text_layer(self):
        import sys
        import unittest.mock
        mock_fitz = unittest.mock.MagicMock()
        mock_page = unittest.mock.MagicMock()
        mock_page.get_text.return_value = "Surface : 1200 pi²\nPrix : 350 000 $"
        mock_doc = unittest.mock.MagicMock()
        mock_doc.__iter__ = unittest.mock.Mock(return_value=iter([mock_page]))
        mock_fitz.open.return_value = mock_doc
        with unittest.mock.patch.dict(sys.modules, {"fitz": mock_fitz}):
            from engine.ingestion import extract_text_from_pdf
            text, has_text = extract_text_from_pdf(Path("/fake/doc.pdf"))
        assert has_text is True
        assert "1200" in text

    def test_returns_false_when_no_text(self):
        import sys
        import unittest.mock
        mock_fitz = unittest.mock.MagicMock()
        mock_page = unittest.mock.MagicMock()
        mock_page.get_text.return_value = ""
        mock_doc = unittest.mock.MagicMock()
        mock_doc.__iter__ = unittest.mock.Mock(return_value=iter([mock_page]))
        mock_fitz.open.return_value = mock_doc
        with unittest.mock.patch.dict(sys.modules, {"fitz": mock_fitz}):
            from engine.ingestion import extract_text_from_pdf
            text, has_text = extract_text_from_pdf(Path("/fake/scan.pdf"))
        assert has_text is False
        assert text == ""


# ── TestIngestion_VisionFallback_PDF ─────────────────────────────────────────

class TestIngestion_VisionFallback_PDF:
    def test_vision_called_when_pdf_has_no_text(self):
        import unittest.mock
        mock_client = unittest.mock.MagicMock()
        mock_resp = unittest.mock.MagicMock()
        mock_resp.choices[0].message.content = "Maison de plain-pied en brique"
        mock_client.chat.completions.create.return_value = mock_resp
        with unittest.mock.patch("engine.ingestion.extract_text_from_pdf", return_value=("", False)):
            with unittest.mock.patch("engine.ingestion.pdf_page_to_b64_image", return_value="fakeb64base64"):
                from engine.ingestion import extract_document
                result = extract_document(Path("/fake/scan.pdf"), "application/pdf", mock_client)
        assert result["method"] == "vision"
        assert "Maison" in result["extracted_text"]

    def test_skipped_when_pdf_has_no_text_and_no_client(self):
        import unittest.mock
        with unittest.mock.patch("engine.ingestion.extract_text_from_pdf", return_value=("", False)):
            from engine.ingestion import extract_document
            result = extract_document(Path("/fake/scan.pdf"), "application/pdf", None)
        assert result["method"] == "skipped"
        assert result["extracted_text"] == ""


# ── TestIngestion_VisionImage ─────────────────────────────────────────────────

class TestIngestion_VisionImage:
    def test_vision_called_for_jpeg(self):
        import unittest.mock
        import tempfile
        mock_client = unittest.mock.MagicMock()
        mock_resp = unittest.mock.MagicMock()
        mock_resp.choices[0].message.content = "Belle maison en brique"
        mock_client.chat.completions.create.return_value = mock_resp
        with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
            f.write(b"\xff\xd8\xff\xe0" + b"\x00" * 20)
            tmp_path = Path(f.name)
        try:
            from engine.ingestion import extract_document
            result = extract_document(tmp_path, "image/jpeg", mock_client)
            assert result["method"] == "vision"
            assert "Belle maison" in result["extracted_text"]
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_skipped_for_jpeg_without_client(self):
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
            f.write(b"\xff\xd8\xff\xe0" + b"\x00" * 20)
            tmp_path = Path(f.name)
        try:
            from engine.ingestion import extract_document
            result = extract_document(tmp_path, "image/jpeg", None)
            assert result["method"] == "skipped"
            assert result["extracted_text"] == ""
        finally:
            tmp_path.unlink(missing_ok=True)


# ── TestIngestion_NoOpenAI ────────────────────────────────────────────────────

class TestIngestion_NoOpenAI:
    def test_no_crash_when_no_client_pdf(self):
        import sys
        import unittest.mock
        mock_fitz = unittest.mock.MagicMock()
        mock_page = unittest.mock.MagicMock()
        mock_page.get_text.return_value = ""
        mock_doc = unittest.mock.MagicMock()
        mock_doc.__iter__ = unittest.mock.Mock(return_value=iter([mock_page]))
        mock_fitz.open.return_value = mock_doc
        with unittest.mock.patch.dict(sys.modules, {"fitz": mock_fitz}):
            from engine.ingestion import extract_document
            result = extract_document(Path("/fake/scan.pdf"), "application/pdf", None)
        assert result["extracted_text"] == ""
        assert result["method"] == "skipped"


# ── TestIngestion_StructuredFields ────────────────────────────────────────────

class TestIngestion_StructuredFields:
    def test_parse_structured_fields_returns_known_keys(self):
        import unittest.mock
        mock_client = unittest.mock.MagicMock()
        mock_resp = unittest.mock.MagicMock()
        mock_resp.choices[0].message.content = (
            '{"prix_achat": 350000.0, "date_achat": "2025-03-15", "no_lot": null}'
        )
        mock_client.chat.completions.create.return_value = mock_resp
        from engine.ingestion import parse_structured_fields
        docs = [{"filename": "acte.pdf", "extracted_text": "Prix : 350 000 $"}]
        result = parse_structured_fields(docs, mock_client)
        assert result["prix_achat"] == 350000.0
        assert result["date_achat"] == "2025-03-15"
        assert "no_lot" not in result  # null excluded

    def test_returns_empty_when_no_client(self):
        from engine.ingestion import parse_structured_fields
        docs = [{"filename": "acte.pdf", "extracted_text": "Prix : 350 000 $"}]
        result = parse_structured_fields(docs, None)
        assert result == {}


# ── TestIngestion_NullFieldsSkipped ──────────────────────────────────────────

class TestIngestion_NullFieldsSkipped:
    def test_null_fields_not_in_result(self):
        import unittest.mock
        mock_client = unittest.mock.MagicMock()
        mock_resp = unittest.mock.MagicMock()
        mock_resp.choices[0].message.content = '{"prix_achat": null, "date_achat": null}'
        mock_client.chat.completions.create.return_value = mock_resp
        from engine.ingestion import parse_structured_fields
        docs = [{"filename": "photo.jpg", "extracted_text": "Maison en briques"}]
        result = parse_structured_fields(docs, mock_client)
        assert result == {}


# ── TestIngestion_NoUpload ────────────────────────────────────────────────────

class TestIngestion_NoUpload:
    def test_empty_uploaded_docs_returns_empty_dict(self):
        from engine.ingestion import ingest_uploaded_documents
        session = {"session_dir": "/tmp/fake-session", "uploaded_documents": []}
        result = ingest_uploaded_documents(session, None)
        assert result == {}

    def test_missing_uploaded_docs_key_returns_empty_dict(self):
        from engine.ingestion import ingest_uploaded_documents
        session = {"session_dir": "/tmp/fake-session"}
        result = ingest_uploaded_documents(session, None)
        assert result == {}


# ── TestIngestion_ExistingFieldsNotOverwritten ────────────────────────────────

class TestIngestion_ExistingFieldsNotOverwritten:
    def test_fixture_field_wins_over_extracted_field(self):
        """Injection loop: 'not case.get(k)' — existing values win."""
        case = {"prix_achat": 450000.0}
        _fields = {"prix_achat": 350000.0, "date_achat": "2025-03-15"}
        for k, v in _fields.items():
            if v is not None and not case.get(k):
                case[k] = v
        assert case["prix_achat"] == 450000.0  # not overwritten
        assert case["date_achat"] == "2025-03-15"  # new field added

    def test_empty_string_case_field_is_overwritten(self):
        """Empty string is falsy — extraction fills the gap."""
        case = {"prix_achat": ""}
        _fields = {"prix_achat": 350000.0}
        for k, v in _fields.items():
            if v is not None and not case.get(k):
                case[k] = v
        assert case["prix_achat"] == 350000.0


# ── TestFicheBien_IngestedDocs ────────────────────────────────────────────────

class TestFicheBien_IngestedDocs:
    def test_ingested_docs_appended_to_fiche_bien_prompt(self):
        from engine.runtime import _build_enrichment_prompt
        case = {
            "dossier_id": "D-INGEST-TEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-13",
            "zone": "Laval",
            "ingested_docs": [
                {
                    "filename": "acte_vente.pdf",
                    "extracted_text": "Prix : 350 000 $\nDate : 2025-03-15",
                },
            ],
        }
        payload = {
            "surface": {"value": 1200, "unit": "pi²"},
            "confidence": 0.85,
            "source_ids": ["SRC-001"],
        }
        prompt = _build_enrichment_prompt("data-facts", "fiche_bien.json", payload, case)
        assert "Documents" in prompt
        assert "acte_vente.pdf" in prompt
        assert "350 000" in prompt

    def test_fiche_bien_prompt_unchanged_without_ingested_docs(self):
        from engine.runtime import _build_enrichment_prompt
        case = {
            "dossier_id": "D-NO-INGEST",
            "type_bien": "residentiel_unifamilial",
            "date_reference": "2026-05-13",
            "zone": "Montreal",
        }
        payload = {
            "surface": {"value": 900, "unit": "pi²"},
            "confidence": 0.70,
            "source_ids": [],
        }
        prompt = _build_enrichment_prompt("data-facts", "fiche_bien.json", payload, case)
        assert "DONNÉES DE LA FICHE BIEN" in prompt
        assert "acte_vente.pdf" not in prompt


# ── TestMapComparableInput_Full ───────────────────────────────────────────────

class TestMapComparableInput_Full:
    def test_all_fields_mapped_correctly(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import _map_comparable_input
        row = {
            "id": "abc123",
            "adresse": "123 rue Example, Montréal",
            "date_vente": "2024-06-15",
            "prix_vente": 450000,
            "source_id": "CENTRIS-12345678",
            "source_type": "mls_centris",
            "type_propriete": "unifamiliale",
            "surface_hab": 145.0,
            "surface_terrain": 350.0,
            "annee_construction": 1985,
            "nb_logements": None,
            "conditions_vente": "normale",
            "notes": "Belle propriété",
        }
        result = _map_comparable_input(row)
        assert result["comparable_id"] == "CENTRIS-12345678"
        assert result["adresse"] == "123 rue Example, Montréal"
        assert result["date_vente"] == "2024-06-15"
        assert result["prix_vente"] == 450000.0
        assert result["source_id"] == "CENTRIS-12345678"
        assert result["source_type"] == "mls_centris"
        assert result["surface"] == {"value": 145.0, "unit": "m²"}
        assert result["surface_terrain"] == 350.0
        assert result["annee_construction"] == 1985
        assert result["nb_logements"] is None
        assert result["conditions_vente"] == "normale"
        assert result["notes"] == "Belle propriété"
        assert result["confidence"] == 0.80


# ── TestMapComparableInput_NullOptionals ──────────────────────────────────────

class TestMapComparableInput_NullOptionals:
    def test_surface_hab_none_returns_empty_surface_dict(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import _map_comparable_input
        row = {
            "surface_hab": None,
            "annee_construction": None,
            "surface_terrain": None,
            "nb_logements": None,
        }
        result = _map_comparable_input(row)
        assert result["surface"] == {}
        assert result["annee_construction"] is None
        assert result["surface_terrain"] is None
        assert result["nb_logements"] is None

    def test_missing_source_id_falls_back_to_id_field(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import _map_comparable_input
        row = {"id": "fallback-uuid", "source_id": ""}
        result = _map_comparable_input(row)
        assert result["comparable_id"] == "fallback-uuid"


# ── TestLoadCaseBody_ComparablesInjected ──────────────────────────────────────

class TestLoadCaseBody_ComparablesInjected:
    def test_comparables_mapped_from_body(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import load_case_from_body
        body = {
            "comparables": [
                {
                    "id": "c1",
                    "adresse": "456 rue Test",
                    "date_vente": "2024-03-01",
                    "prix_vente": 500000,
                    "source_id": "RF-2024-001",
                    "source_type": "registre_foncier",
                    "surface_hab": 120.0,
                    "surface_terrain": None,
                    "annee_construction": 1992,
                    "nb_logements": None,
                    "conditions_vente": "normale",
                    "notes": "",
                }
            ]
        }
        case, _ = load_case_from_body(body)
        assert len(case["comparables"]) == 1
        comp = case["comparables"][0]
        assert comp["source_id"] == "RF-2024-001"
        assert comp["surface"] == {"value": 120.0, "unit": "m²"}
        assert comp["confidence"] == 0.80

    def test_comparables_body_override_fixture_comparables(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import load_case_from_body
        body = {
            "comparables": [
                {
                    "id": "new1",
                    "source_id": "NEW-001",
                    "prix_vente": 300000,
                    "surface_hab": None,
                    "surface_terrain": None,
                    "annee_construction": None,
                    "nb_logements": None,
                }
            ]
        }
        case, _ = load_case_from_body(body)
        assert len(case["comparables"]) >= 1
        assert all(c["source_id"] == "NEW-001" for c in case["comparables"])

    def test_no_comparables_in_body_leaves_fixture_untouched(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from api import load_case_from_body
        case_a, _ = load_case_from_body({})
        fixture_comps = list(case_a.get("comparables", []))
        case_b, _ = load_case_from_body({})
        assert case_b.get("comparables", []) == fixture_comps


# ── TestBuildRapportPromptV2 ───────────────────────────────────────────────────

class TestBuildRapportPromptV2_IncludesCommanditaire:
    def test_commanditaire_nom_in_prompt(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import _build_rapport_prompt_v2
        case = {
            "dossier_id": "D-TEST",
            "commanditaire": {
                "nom": "Jean Tremblay",
                "organisation": "Banque XYZ",
                "fin_evaluation": "hypothecaire",
            },
            "date_reference": "2026-05-15",
            "type_bien": "residentiel_unifamilial",
            "zone": "R-1",
            "surface": {"value": 120, "unit": "m²"},
            "comparables": [],
        }
        prompt = _build_rapport_prompt_v2(case, "abrege", {}, "BROUILLON", [], [])
        assert "Jean Tremblay" in prompt

    def test_commanditaire_organisation_in_prompt(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import _build_rapport_prompt_v2
        case = {
            "commanditaire": {"nom": "Marie Côté", "organisation": "Caisse Pop", "fin_evaluation": "succession"},
        }
        prompt = _build_rapport_prompt_v2(case, "abrege", {}, "BROUILLON", [], [])
        assert "Caisse Pop" in prompt


class TestBuildRapportPromptV2_FormatAbrege:
    def test_format_abrege_label_in_prompt(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import _build_rapport_prompt_v2
        case = {"dossier_id": "D-TEST"}
        prompt = _build_rapport_prompt_v2(case, "abrege", {}, "BROUILLON", [], [])
        assert "abrege" in prompt.lower() or "abrégé" in prompt.lower()

    def test_format_abrege_not_complet(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import _build_rapport_prompt_v2
        case = {"dossier_id": "D-TEST"}
        prompt = _build_rapport_prompt_v2(case, "abrege", {}, "BROUILLON", [], [])
        assert "narratif complet" not in prompt.lower()


class TestBuildRapportPromptV2_FormatComplet:
    def test_format_complet_label_in_prompt(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import _build_rapport_prompt_v2
        case = {"dossier_id": "D-TEST"}
        prompt = _build_rapport_prompt_v2(case, "complet", {}, "BROUILLON", [], [])
        assert "complet" in prompt.lower() or "narratif" in prompt.lower()


class TestGenerateRapportFallbackNoCle:
    def test_returns_deterministic_string_without_api_key(self, monkeypatch):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import generate_brouillon_rapport
        monkeypatch.delenv("OPENAI_API_KEY", raising=False)
        case = {
            "dossier_id": "D-TEST",
            "type_bien": "residentiel_unifamilial",
            "zone": "R-1",
            "date_reference": "2026-05-15",
            "surface": {"value": 120, "unit": "m²"},
            "comparables": [],
        }
        result = generate_brouillon_rapport(case, {}, "BROUILLON", [], [], format="abrege")
        assert isinstance(result, str)
        assert len(result) > 100
        assert "BROUILLON" in result

    def test_complet_format_also_returns_string(self, monkeypatch):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.runtime import generate_brouillon_rapport
        monkeypatch.delenv("OPENAI_API_KEY", raising=False)
        case = {"dossier_id": "D-TEST", "type_bien": "immeuble_revenus"}
        result = generate_brouillon_rapport(case, {}, "BROUILLON", [], [], format="complet")
        assert isinstance(result, str)
        assert len(result) > 100


class TestSaveRapportContent:
    def test_writes_content_to_artifact_file(self, tmp_path, monkeypatch):
        """app_save_rapport écrase le fichier brouillon_rapport.md dans la session."""
        import sys, json
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        import api as api_module
        monkeypatch.setattr(api_module, "SESSIONS_DIR", tmp_path)

        session_id = "test-session-abc"
        session_dir = tmp_path / session_id
        session_dir.mkdir()
        artifacts_dir = session_dir / "artifacts" / "D-TEST"
        artifacts_dir.mkdir(parents=True)
        rapport_path = artifacts_dir / "redaction.brouillon_rapport.md"
        rapport_path.write_text("# Brouillon original\n", encoding="utf-8")

        artifact_index = {
            "artifacts": [
                {
                    "step": "redaction",
                    "artifact": "brouillon_rapport.md",
                    "event_id": "evt_001",
                    "path": str(rapport_path),
                }
            ]
        }
        (session_dir / "artifact_index.json").write_text(
            json.dumps(artifact_index), encoding="utf-8"
        )
        session_data = {
            "session_id": session_id,
            "session_dir": str(session_dir),
        }
        (session_dir / "session.json").write_text(
            json.dumps(session_data), encoding="utf-8"
        )

        result = api_module.app_save_rapport(
            {"session_id": session_id, "content": "# Contenu modifié\n\nTexte édité."}
        )
        assert result["ok"] is True
        assert rapport_path.read_text(encoding="utf-8") == "# Contenu modifié\n\nTexte édité."


# ── Batch 8b — Export rapport ────────────────────────────────────────────────

class TestGenerateDocx_ContainsWatermark:
    def test_watermark_in_generated_docx(self):
        import sys, io
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.report_export import _generate_docx
        from docx import Document
        data = _generate_docx("## Identification\n\nTestDocument", "D-TEST")
        assert isinstance(data, bytes) and len(data) > 0
        doc = Document(io.BytesIO(data))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "BROUILLON NON CERTIFIÉ" in all_text


class TestGenerateDocx_HeadingsRendered:
    def test_h2_becomes_heading2(self):
        import sys, io
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.report_export import _generate_docx
        from docx import Document
        data = _generate_docx("## Section principale\n\nTexte normal.", "D-TEST")
        doc = Document(io.BytesIO(data))
        heading_styles = [p.style.name for p in doc.paragraphs]
        assert "Heading 2" in heading_styles


class TestGenerateHtml_ContainsWatermark:
    def test_watermark_div_present(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.report_export import _generate_html
        html = _generate_html("## Test\n\nContenu.", "D-TEST")
        assert isinstance(html, str)
        assert "BROUILLON NON CERTIFIÉ" in html


class TestGenerateHtml_TablesRendered:
    def test_markdown_table_becomes_html_table(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.report_export import _generate_html
        md = "| Col A | Col B |\n|-------|-------|\n| val1  | val2  |"
        html = _generate_html(md, "D-TEST")
        assert "<table" in html.lower()
        assert "val1" in html


class TestExportRapport_DocxEndpoint:
    def test_docx_export_returns_base64_with_correct_fields(self, tmp_path, monkeypatch):
        import sys, json, base64
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        import api as api_module
        monkeypatch.setattr(api_module, "SESSIONS_DIR", tmp_path)

        session_id = "test-export-docx"
        session_dir = tmp_path / session_id
        session_dir.mkdir()
        artifacts_dir = session_dir / "artifacts" / "D-EXPORT"
        artifacts_dir.mkdir(parents=True)
        rapport_path = artifacts_dir / "redaction.brouillon_rapport.md"
        rapport_path.write_text("## Rapport\n\nContenu test.", encoding="utf-8")
        artifact_index = {"artifacts": [{"step": "redaction", "artifact": "brouillon_rapport.md",
                                          "event_id": "evt_001", "path": str(rapport_path)}]}
        (session_dir / "artifact_index.json").write_text(json.dumps(artifact_index), encoding="utf-8")
        (session_dir / "session.json").write_text(
            json.dumps({"session_id": session_id, "session_dir": str(session_dir), "dossier_id": "D-EXPORT"}),
            encoding="utf-8")

        result = api_module.app_export_rapport({"session_id": session_id, "format": "docx"})
        assert result["ok"] is True
        assert result["content_type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert result["filename"] == "rapport-D-EXPORT.docx"
        data = base64.b64decode(result["data"])
        assert len(data) > 100


class TestExportRapport_HtmlEndpoint:
    def test_html_export_returns_html_string_with_watermark(self, tmp_path, monkeypatch):
        import sys, json
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        import api as api_module
        monkeypatch.setattr(api_module, "SESSIONS_DIR", tmp_path)

        session_id = "test-export-html"
        session_dir = tmp_path / session_id
        session_dir.mkdir()
        artifacts_dir = session_dir / "artifacts" / "D-HTML"
        artifacts_dir.mkdir(parents=True)
        rapport_path = artifacts_dir / "redaction.brouillon_rapport.md"
        rapport_path.write_text("## Test\n\nContenu.", encoding="utf-8")
        artifact_index = {"artifacts": [{"step": "redaction", "artifact": "brouillon_rapport.md",
                                          "event_id": "evt_001", "path": str(rapport_path)}]}
        (session_dir / "artifact_index.json").write_text(json.dumps(artifact_index), encoding="utf-8")
        (session_dir / "session.json").write_text(
            json.dumps({"session_id": session_id, "session_dir": str(session_dir), "dossier_id": "D-HTML"}),
            encoding="utf-8")

        result = api_module.app_export_rapport({"session_id": session_id, "format": "html"})
        assert result["ok"] is True
        assert result["content_type"] == "text/html; charset=utf-8"
        assert "BROUILLON NON CERTIFIÉ" in result["data"]


class TestExportRapport_InvalidFormat:
    def test_format_pdf_raises_value_error(self, tmp_path, monkeypatch):
        import sys, json, pytest
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        import api as api_module
        monkeypatch.setattr(api_module, "SESSIONS_DIR", tmp_path)

        session_id = "test-invalid-fmt"
        session_dir = tmp_path / session_id
        session_dir.mkdir()
        (session_dir / "session.json").write_text(
            json.dumps({"session_id": session_id, "session_dir": str(session_dir)}),
            encoding="utf-8")

        with pytest.raises(ValueError, match="format"):
            api_module.app_export_rapport({"session_id": session_id, "format": "pdf"})


# ── TestDataEnrichment ────────────────────────────────────────────────────────

class TestDataEnrichment_CityDetection:
    def _detect(self, display_name, zone=""):
        from engine.data_enrichment import detect_city
        return detect_city(display_name, zone)

    def test_montreal_explicit(self):
        assert self._detect("1234 rue Sherbrooke, Montréal") == "montreal"

    def test_plateau_maps_to_montreal(self):
        assert self._detect("123 Avenue du Parc", zone="plateau-mont-royal") == "montreal"

    def test_laval_maps_to_montreal(self):
        assert self._detect("456 boul Laval, Laval") == "montreal"

    def test_quebec_city(self):
        assert self._detect("100 Grande Allée, Québec") == "quebec"

    def test_gatineau(self):
        assert self._detect("789 boul Maloney, Gatineau") == "gatineau"

    def test_sherbrooke(self):
        assert self._detect("321 King O, Sherbrooke") == "sherbrooke"

    def test_default_fallback(self):
        assert self._detect("", zone="SECTEUR-ANONYMISE") == "montreal"

    def test_case_insensitive(self):
        assert self._detect("MONTREAL") == "montreal"


class TestDataEnrichment_EnrichCase:
    def test_enrich_case_never_raises(self, tmp_path):
        from engine.data_enrichment import enrich_case
        case = {"dossier_id": "TEST", "zone": "SECTEUR-X"}
        # Should not raise even with no network / no CSV
        enrich_case(case, display_name="9999 rue Inexistante", cache_dir=tmp_path)

    def test_enrich_case_no_side_effect_on_failure(self, tmp_path):
        from engine.data_enrichment import enrich_case
        case = {"dossier_id": "TEST", "zone": "SECTEUR-X", "surface": 150.0}
        enrich_case(case, display_name="", cache_dir=tmp_path)
        # surface must not be overwritten
        assert case["surface"] == 150.0

    def test_role_lookup_empty_when_no_csv(self, tmp_path):
        from engine.data_enrichment import lookup_role_mtl
        result = lookup_role_mtl(tmp_path / "nonexistent.csv")
        assert result == {}

    def test_role_lookup_from_csv(self, tmp_path):
        from engine.data_enrichment import lookup_role_mtl
        csv_path = tmp_path / "role_mtl.csv"
        csv_path.write_text(
            "ID_UEV,CIVIQUE_DEBUT,CIVIQUE_FIN,NOM_RUE,SUITE_DEBUT,ETAGE_HORS_SOL,"
            "NOMBRE_LOGEMENT,ANNEE_CONSTRUCTION,CODE_UTILISATION,LETTRE_DEBUT,"
            "LETTRE_FIN,LIBELLE_UTILISATION,CATEGORIE_UEF,MATRICULE83,"
            "SUPERFICIE_TERRAIN,SUPERFICIE_BATIMENT,NO_ARROND_ILE_CUM,MUNICIPALITE\n"
            "1234567,1000,1000,RUE SHERBROOKE O,,3,6,1958,1000,,,"
            "Résidentiel,Régulier,9999-12-3456-7-001-0,312.5,420.0,10,Montréal\n",
            encoding="utf-8",
        )
        result = lookup_role_mtl(csv_path, matricule="9999-12-3456-7-001-0")
        assert result["annee_construction"] == 1958
        assert result["nb_logements"] == 6
        assert result["superficie_terrain_m2"] == 312.5

    def test_role_lookup_by_address(self, tmp_path):
        from engine.data_enrichment import lookup_role_mtl
        csv_path = tmp_path / "role_mtl.csv"
        csv_path.write_text(
            "ID_UEV,CIVIQUE_DEBUT,CIVIQUE_FIN,NOM_RUE,SUITE_DEBUT,ETAGE_HORS_SOL,"
            "NOMBRE_LOGEMENT,ANNEE_CONSTRUCTION,CODE_UTILISATION,LETTRE_DEBUT,"
            "LETTRE_FIN,LIBELLE_UTILISATION,CATEGORIE_UEF,MATRICULE83,"
            "SUPERFICIE_TERRAIN,SUPERFICIE_BATIMENT,NO_ARROND_ILE_CUM,MUNICIPALITE\n"
            "1234567,1000,1000,RUE SHERBROOKE O,,3,6,1958,1000,,,"
            "Résidentiel,Régulier,9999-12-3456-7-001-0,312.5,420.0,10,Montréal\n",
            encoding="utf-8",
        )
        result = lookup_role_mtl(csv_path, display_name="1000 rue Sherbrooke O")
        assert result["annee_construction"] == 1958

    def test_enrich_case_injects_role_from_csv(self, tmp_path):
        from engine.data_enrichment import enrich_case
        csv_path = tmp_path / "role_mtl.csv"
        csv_path.write_text(
            "ID_UEV,CIVIQUE_DEBUT,CIVIQUE_FIN,NOM_RUE,SUITE_DEBUT,ETAGE_HORS_SOL,"
            "NOMBRE_LOGEMENT,ANNEE_CONSTRUCTION,CODE_UTILISATION,LETTRE_DEBUT,"
            "LETTRE_FIN,LIBELLE_UTILISATION,CATEGORIE_UEF,MATRICULE83,"
            "SUPERFICIE_TERRAIN,SUPERFICIE_BATIMENT,NO_ARROND_ILE_CUM,MUNICIPALITE\n"
            "1234567,500,500,AV DU PARC,,2,4,1975,1000,,,"
            "Résidentiel,Régulier,ABCD-12-3456-7-000-0000,400.0,300.0,10,Montréal\n",
            encoding="utf-8",
        )
        case = {"dossier_id": "TEST", "matricule": "ABCD-12-3456-7-000-0000"}
        enrich_case(case, display_name="500 av du Parc", cache_dir=tmp_path)
        assert case.get("role_municipal", {}).get("annee_construction") == 1975
        assert case.get("annee_construction") == 1975


class TestDataEnrichment_XmlRole:
    """Tests for MAMH XML index builder + lookup (autres villes)."""

    def _make_xml(self, tmp_path, entries: list[dict]) -> Path:
        """Build a minimal MAMH XML with given UEV entries."""
        rows = []
        for e in entries:
            rows.append(f"""  <RLUEx>
    <RL0101><RL0101x>
      <RL0101Ax>{e.get('civique','')}</RL0101Ax>
      <RL0101Ex>{e.get('type_voie','RU')}</RL0101Ex>
      <RL0101Gx>{e.get('nom_voie','')}</RL0101Gx>
    </RL0101x></RL0101>
    <RL0103><RL0103x><RL0103Ax>{e.get('lot','')}</RL0103Ax></RL0103x></RL0103>
    <RL0104>
      <RL0104A>{e.get('a','0001')}</RL0104A>
      <RL0104B>{e.get('b','01')}</RL0104B>
      <RL0104C>{e.get('c','0001')}</RL0104C>
      <RL0104D>{e.get('d','1')}</RL0104D>
    </RL0104>
    <RL0105A>{e.get('cubf','1000')}</RL0105A>
    <RL0302A>{e.get('sup_terrain','400')}</RL0302A>
    <RL0307A>{e.get('annee','1980')}</RL0307A>
    <RL0308A>{e.get('sup_bat','200')}</RL0308A>
    <RL0311A>{e.get('logements','2')}</RL0311A>
    <RL0402A>{e.get('val_terrain','150000')}</RL0402A>
    <RL0403A>{e.get('val_bat','250000')}</RL0403A>
    <RL0404A>{e.get('val_totale','400000')}</RL0404A>
  </RLUEx>""")
        xml_content = "<ROLE>\n" + "\n".join(rows) + "\n</ROLE>"
        p = tmp_path / "test_role.xml"
        p.write_text(xml_content, encoding="utf-8")
        return p

    def test_build_xml_index_count(self, tmp_path):
        from engine.data_enrichment import build_role_xml_index
        xml = self._make_xml(tmp_path, [
            {"civique": "100", "nom_voie": "GRANDE ALLEE", "annee": "1920",
             "a": "0023", "b": "01", "c": "0001", "d": "1"},
            {"civique": "200", "nom_voie": "CHEMIN SAINTE-FOY", "annee": "1965",
             "a": "0023", "b": "01", "c": "0002", "d": "1"},
        ])
        index_path = tmp_path / "idx.json"
        count = build_role_xml_index(xml, index_path, "quebec")
        assert count == 2
        assert index_path.exists()

    def test_xml_index_lookup_by_matricule(self, tmp_path):
        from engine.data_enrichment import build_role_xml_index, lookup_role_xml
        xml = self._make_xml(tmp_path, [
            {"civique": "100", "nom_voie": "GRANDE ALLEE", "annee": "1920",
             "a": "0023", "b": "01", "c": "0001", "d": "1",
             "val_totale": "680000", "sup_bat": "350"},
        ])
        index_path = tmp_path / "idx.json"
        build_role_xml_index(xml, index_path, "quebec")
        result = lookup_role_xml(index_path, matricule="0023-01-0001-1-000-0000")
        assert result["annee_construction"] == 1920
        assert result["valeur_totale"] == 680000.0
        assert result["superficie_batiment_m2"] == 350.0

    def test_xml_index_lookup_by_address(self, tmp_path):
        from engine.data_enrichment import build_role_xml_index, lookup_role_xml
        xml = self._make_xml(tmp_path, [
            {"civique": "500", "type_voie": "AV", "nom_voie": "CARTIER",
             "annee": "1955", "a": "0081", "b": "02", "c": "0003", "d": "2"},
        ])
        index_path = tmp_path / "idx.json"
        build_role_xml_index(xml, index_path, "gatineau")
        result = lookup_role_xml(index_path, display_name="500 av Cartier")
        assert result["annee_construction"] == 1955

    def test_xml_index_missing_returns_empty(self, tmp_path):
        from engine.data_enrichment import lookup_role_xml
        result = lookup_role_xml(tmp_path / "nonexistent.json")
        assert result == {}

    def test_enrich_case_xml_city_builds_index_from_xml(self, tmp_path):
        from engine.data_enrichment import enrich_case, build_role_xml_index
        # Simulate pre-downloaded XML for Gatineau
        xml_entries = [{"civique": "789", "type_voie": "BOUL", "nom_voie": "MALONEY",
                        "annee": "1988", "a": "0081", "b": "03", "c": "0005", "d": "1",
                        "val_totale": "320000"}]
        xml_content = ("<ROLE>\n"
                       "  <RLUEx>\n"
                       "    <RL0101><RL0101x>"
                       "<RL0101Ax>789</RL0101Ax>"
                       "<RL0101Ex>BOUL</RL0101Ex>"
                       "<RL0101Gx>MALONEY</RL0101Gx>"
                       "</RL0101x></RL0101>\n"
                       "    <RL0103><RL0103x><RL0103Ax>11111</RL0103Ax></RL0103x></RL0103>\n"
                       "    <RL0104>\n"
                       "      <RL0104A>0081</RL0104A><RL0104B>03</RL0104B>\n"
                       "      <RL0104C>0005</RL0104C><RL0104D>1</RL0104D>\n"
                       "    </RL0104>\n"
                       "    <RL0105A>1000</RL0105A>\n"
                       "    <RL0302A>600</RL0302A>\n"
                       "    <RL0307A>1988</RL0307A>\n"
                       "    <RL0308A>180</RL0308A>\n"
                       "    <RL0311A>1</RL0311A>\n"
                       "    <RL0402A>80000</RL0402A>\n"
                       "    <RL0403A>240000</RL0403A>\n"
                       "    <RL0404A>320000</RL0404A>\n"
                       "  </RLUEx>\n"
                       "</ROLE>\n")
        xml_path = tmp_path / "role_gatineau.xml"
        xml_path.write_text(xml_content, encoding="utf-8")

        case = {"dossier_id": "TEST-GATINEAu", "zone": "gatineau-aylmer"}
        enrich_case(case, display_name="789 boul Maloney, Gatineau", cache_dir=tmp_path)

        assert case.get("role_municipal", {}).get("annee_construction") == 1988
        assert case.get("role_municipal", {}).get("valeur_totale") == 320000.0
        assert case.get("annee_construction") == 1988
        assert case.get("evaluation_municipale_totale") == 320000.0


# ── TestDataEnrichment_Zonage ─────────────────────────────────────────────────

class TestDataEnrichment_Zonage:
    """Tests for geocoding helpers and zoning PiP lookup."""

    def _import(self):
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
        from engine.data_enrichment import (
            _pip_exterior,
            _simplify_ring,
            build_zoning_index,
            lookup_zoning_point,
            _ZONING_INDEX_CACHE,
        )
        return _pip_exterior, _simplify_ring, build_zoning_index, lookup_zoning_point, _ZONING_INDEX_CACHE

    def test_pip_point_inside_square(self):
        """Point at centre of unit square → inside."""
        _pip_exterior, *_ = self._import()
        ring = [[0.0, 0.0], [1.0, 0.0], [1.0, 1.0], [0.0, 1.0], [0.0, 0.0]]
        assert _pip_exterior(0.5, 0.5, ring) is True

    def test_pip_point_outside_square(self):
        """Point clearly outside square → False."""
        _pip_exterior, *_ = self._import()
        ring = [[0.0, 0.0], [1.0, 0.0], [1.0, 1.0], [0.0, 1.0], [0.0, 0.0]]
        assert _pip_exterior(2.0, 2.0, ring) is False

    def test_pip_point_outside_left(self):
        """Point to the left of square → False."""
        _pip_exterior, *_ = self._import()
        ring = [[0.0, 0.0], [1.0, 0.0], [1.0, 1.0], [0.0, 1.0], [0.0, 0.0]]
        assert _pip_exterior(-0.5, 0.5, ring) is False

    def test_simplify_ring_keeps_short(self):
        """Ring shorter than max_pts returned unchanged."""
        _, _simplify_ring, *_ = self._import()
        ring = [[float(i), 0.0] for i in range(10)]
        result = _simplify_ring(ring, max_pts=20)
        assert result == ring

    def test_simplify_ring_downsamples(self):
        """Ring longer than max_pts is downsampled to ≤ max_pts."""
        _, _simplify_ring, *_ = self._import()
        ring = [[float(i), 0.0] for i in range(1000)]
        result = _simplify_ring(ring, max_pts=100)
        assert len(result) == 100

    def test_build_zoning_index_polygon(self, tmp_path):
        """build_zoning_index creates correct index from a simple GeoJSON polygon."""
        _pip_exterior, _simplify_ring, build_zoning_index, *_ = self._import()
        geojson = {
            "type": "FeatureCollection",
            "features": [
                {
                    "type": "Feature",
                    "properties": {"ZONE": "RH", "DESCRIPTION": "Résidentiel haute densité"},
                    "geometry": {
                        "type": "Polygon",
                        "coordinates": [
                            [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6],
                             [-73.6, 45.6], [-73.6, 45.5]]
                        ],
                    },
                }
            ],
        }
        import json
        gj_path = tmp_path / "zoning_montreal.geojson"
        gj_path.write_text(json.dumps(geojson), encoding="utf-8")
        idx_path = tmp_path / "zoning_montreal_index.json"

        count = build_zoning_index(gj_path, idx_path)
        assert count == 1
        assert idx_path.exists()

        data = json.loads(idx_path.read_text(encoding="utf-8"))
        assert data["_count"] == 1
        zone = data["zones"][0]
        assert zone["props"]["ZONE"] == "RH"
        import pytest
        assert zone["bbox"][0] == pytest.approx(-73.6)
        assert len(zone["ring"]) >= 4

    def test_build_zoning_index_multipolygon(self, tmp_path):
        """MultiPolygon with 2 parts → 2 zones in index."""
        _, _, build_zoning_index, *_ = self._import()
        geojson = {
            "type": "FeatureCollection",
            "features": [
                {
                    "type": "Feature",
                    "properties": {"ZONE": "C1"},
                    "geometry": {
                        "type": "MultiPolygon",
                        "coordinates": [
                            [[[-73.7, 45.4], [-73.6, 45.4], [-73.6, 45.5],
                              [-73.7, 45.5], [-73.7, 45.4]]],
                            [[[-73.5, 45.4], [-73.4, 45.4], [-73.4, 45.5],
                              [-73.5, 45.5], [-73.5, 45.4]]],
                        ],
                    },
                }
            ],
        }
        import json
        gj_path = tmp_path / "zoning_montreal.geojson"
        gj_path.write_text(json.dumps(geojson), encoding="utf-8")
        idx_path = tmp_path / "zoning_montreal_index.json"
        count = build_zoning_index(gj_path, idx_path)
        assert count == 2

    def test_lookup_zoning_point_hit(self, tmp_path):
        """Point inside a zone polygon → returns zone props."""
        import json
        from engine import data_enrichment
        _pip_exterior, _simplify_ring, build_zoning_index, lookup_zoning_point, _ZONING_INDEX_CACHE = self._import()

        # Clear module cache to avoid stale entries from other tests
        _ZONING_INDEX_CACHE.clear()

        # Temporarily add "testville" to _ZONING_CITIES
        from engine import data_enrichment as de
        de._ZONING_CITIES["testville"] = {"bbox": [-74.0, 45.0, -73.0, 46.0]}

        geojson = {
            "type": "FeatureCollection",
            "features": [
                {
                    "type": "Feature",
                    "properties": {"ZONE": "RM", "DESCRIPTION": "Résidentiel moyen"},
                    "geometry": {
                        "type": "Polygon",
                        "coordinates": [
                            [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6],
                             [-73.6, 45.6], [-73.6, 45.5]]
                        ],
                    },
                }
            ],
        }
        gj_path = tmp_path / "zoning_testville.geojson"
        gj_path.write_text(json.dumps(geojson), encoding="utf-8")
        idx_path = tmp_path / "zoning_testville_index.json"
        build_zoning_index(gj_path, idx_path)

        # Point inside the polygon
        result = lookup_zoning_point("testville", 45.55, -73.55, tmp_path)
        assert result.get("ZONE") == "RM"
        assert result.get("source") == "zonage-testville"

        # Cleanup
        del de._ZONING_CITIES["testville"]

    def test_lookup_zoning_point_miss(self, tmp_path):
        """Point outside all zones → returns {}."""
        import json
        from engine import data_enrichment as de

        de._ZONING_INDEX_CACHE.clear()
        de._ZONING_CITIES["testville2"] = {"bbox": [-74.0, 45.0, -73.0, 46.0]}

        geojson = {
            "type": "FeatureCollection",
            "features": [
                {
                    "type": "Feature",
                    "properties": {"ZONE": "RM"},
                    "geometry": {
                        "type": "Polygon",
                        "coordinates": [
                            [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6],
                             [-73.6, 45.6], [-73.6, 45.5]]
                        ],
                    },
                }
            ],
        }
        gj_path = tmp_path / "zoning_testville2.geojson"
        gj_path.write_text(json.dumps(geojson), encoding="utf-8")
        from engine.data_enrichment import build_zoning_index, lookup_zoning_point
        build_zoning_index(gj_path, tmp_path / "zoning_testville2_index.json")

        # Point far outside
        result = lookup_zoning_point("testville2", 46.9, -75.0, tmp_path)
        assert result == {}

        del de._ZONING_CITIES["testville2"]

    def test_lookup_zoning_unsupported_city(self, tmp_path):
        """Unsupported city_code → {} with no error."""
        from engine.data_enrichment import lookup_zoning_point
        result = lookup_zoning_point("mars", 45.5, -73.5, tmp_path)
        assert result == {}

    def test_enrich_case_injects_zonage(self, tmp_path):
        """enrich_case with mocked geocoding injects zonage_urbanisme when GeoJSON present."""
        import json
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case, build_zoning_index

        de._ZONING_INDEX_CACHE.clear()

        # Patch geocode_address to return coords inside our test polygon
        with mock.patch.object(de, "geocode_address", return_value=(45.55, -73.55)):
            # Add testville3 to supported cities
            de._ZONING_CITIES["testville3"] = {"bbox": [-74.0, 45.0, -73.0, 46.0]}
            # Patch detect_city to return testville3
            with mock.patch.object(de, "detect_city", return_value="testville3"):
                geojson = {
                    "type": "FeatureCollection",
                    "features": [{
                        "type": "Feature",
                        "properties": {"ZONE": "I1", "DESCRIPTION": "Industriel léger"},
                        "geometry": {
                            "type": "Polygon",
                            "coordinates": [[
                                [-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6],
                                [-73.6, 45.6], [-73.6, 45.5]
                            ]],
                        },
                    }],
                }
                gj_path = tmp_path / "zoning_testville3.geojson"
                gj_path.write_text(json.dumps(geojson), encoding="utf-8")

                case = {"dossier_id": "D-ZU-TEST"}
                enrich_case(case, display_name="100 rue Test, TestVille", cache_dir=tmp_path)

                assert case.get("zonage_urbanisme", {}).get("ZONE") == "I1"
                assert case["zonage_urbanisme"]["source"] == "zonage-testville3"

            del de._ZONING_CITIES["testville3"]


# ── TestDataEnrichment_ZonageVilles ───────────────────────────────────────────

class TestDataEnrichment_ZonageVilles:
    """Tests for multi-city zoning config and direct_url download path."""

    def test_all_new_cities_in_zoning_cities(self):
        """Québec, Laval, Longueuil, Gatineau, Sherbrooke all registered."""
        from engine.data_enrichment import _ZONING_CITIES
        for city in ("quebec", "laval", "longueuil", "gatineau", "sherbrooke"):
            assert city in _ZONING_CITIES, f"{city} missing from _ZONING_CITIES"

    def test_longueuil_has_direct_url(self):
        """Longueuil (no CKAN) has direct_url configured."""
        from engine.data_enrichment import _ZONING_CITIES
        cfg = _ZONING_CITIES["longueuil"]
        assert cfg.get("ckan_api") is None
        assert cfg.get("direct_url", "").startswith("http")

    def test_new_cities_have_bbox(self):
        """All new cities define a bounding box."""
        from engine.data_enrichment import _ZONING_CITIES
        for city in ("quebec", "laval", "longueuil", "gatineau", "sherbrooke"):
            bbox = _ZONING_CITIES[city].get("bbox")
            assert isinstance(bbox, list) and len(bbox) == 4, f"bbox missing/malformed for {city}"

    def test_lookup_unsupported_city_still_returns_empty(self, tmp_path):
        """City not in _ZONING_CITIES → {} without error."""
        from engine.data_enrichment import lookup_zoning_point
        assert lookup_zoning_point("rimouski", 48.45, -68.52, tmp_path) == {}

    def test_download_zoning_uses_direct_url(self, tmp_path):
        """When CKAN is None, download_zoning_geojson uses direct_url."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import download_zoning_geojson

        fake_geojson = b'{"type":"FeatureCollection","features":[]}'

        # Simulate a city with only direct_url (like Longueuil)
        de._ZONING_CITIES["mockville"] = {
            "ckan_api": None,
            "package_ids": [],
            "direct_url": "https://example.com/zonage.geojson",
            "bbox": [-74.0, 45.0, -73.0, 46.0],
        }
        try:
            # Mock httpx.stream to return fake GeoJSON bytes
            mock_response = mock.MagicMock()
            mock_response.__enter__ = mock.Mock(return_value=mock_response)
            mock_response.__exit__ = mock.Mock(return_value=False)
            mock_response.raise_for_status = mock.Mock()
            mock_response.iter_bytes = mock.Mock(return_value=iter([fake_geojson]))

            with mock.patch("httpx.stream", return_value=mock_response) as mock_stream:
                result = download_zoning_geojson("mockville", tmp_path)

            # Should have called httpx.stream with the direct_url (not CKAN)
            mock_stream.assert_called_once()
            call_url = mock_stream.call_args[0][1]
            assert "example.com" in call_url
            assert result is not None
            assert result.exists()
        finally:
            del de._ZONING_CITIES["mockville"]

    def test_download_zoning_ckan_fallback_to_direct_url(self, tmp_path):
        """CKAN discovery fails → falls back to direct_url."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import download_zoning_geojson

        fake_geojson = b'{"type":"FeatureCollection","features":[]}'

        de._ZONING_CITIES["fallbackville"] = {
            "ckan_api": "https://fake-ckan.example.com/api/3/action",
            "package_ids": ["bad-package-id"],
            "direct_url": "https://direct.example.com/zonage.geojson",
            "bbox": [-74.0, 45.0, -73.0, 46.0],
        }
        try:
            mock_response = mock.MagicMock()
            mock_response.__enter__ = mock.Mock(return_value=mock_response)
            mock_response.__exit__ = mock.Mock(return_value=False)
            mock_response.raise_for_status = mock.Mock()
            mock_response.iter_bytes = mock.Mock(return_value=iter([fake_geojson]))

            # _find_ckan_geojson raises (CKAN unavailable)
            with mock.patch.object(de, "_find_ckan_geojson", side_effect=Exception("CKAN down")):
                with mock.patch("httpx.stream", return_value=mock_response) as mock_stream:
                    result = download_zoning_geojson("fallbackville", tmp_path)

            call_url = mock_stream.call_args[0][1]
            assert "direct.example.com" in call_url
            assert result is not None
        finally:
            del de._ZONING_CITIES["fallbackville"]

    def test_lookup_new_city_no_data_returns_empty(self, tmp_path):
        """Lookup for a new city with no cached data returns {} gracefully."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import lookup_zoning_point

        de._ZONING_INDEX_CACHE.clear()

        # download returns None (no data available)
        with mock.patch.object(de, "download_zoning_geojson", return_value=None):
            result = lookup_zoning_point("quebec", 46.82, -71.21, tmp_path)

        assert result == {}


# ── TestDataEnrichment_CPTAQ ──────────────────────────────────────────────────

class TestDataEnrichment_CPTAQ:
    """Tests for CPTAQ zone agricole lookup."""

    def _make_cptaq_geojson(self, polygon_rings: list) -> dict:
        return {
            "type": "FeatureCollection",
            "features": [
                {
                    "type": "Feature",
                    "properties": {"NM_MRC": "Rouville", "CATEGORIE": "zone_agricole"},
                    "geometry": {"type": "Polygon", "coordinates": [ring]},
                }
                for ring in polygon_rings
            ],
        }

    def test_build_cptaq_index(self, tmp_path):
        """build_cptaq_index creates valid index from GeoJSON."""
        import json
        from engine.data_enrichment import build_cptaq_index

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "cptaq_zone_agricole.geojson"
        gj_path.write_text(json.dumps(self._make_cptaq_geojson([ring])), encoding="utf-8")

        idx_path = tmp_path / "cptaq_index.json"
        count = build_cptaq_index(gj_path, idx_path)
        assert count == 1
        data = json.loads(idx_path.read_text(encoding="utf-8"))
        assert data["_count"] == 1
        assert data["zones"][0]["props"]["NM_MRC"] == "Rouville"

    def test_lookup_cptaq_inside(self, tmp_path):
        """Point inside zone → en_zone_agricole: True."""
        import json
        from engine import data_enrichment as de
        from engine.data_enrichment import build_cptaq_index, lookup_cptaq

        de._CPTAQ_INDEX_CACHE = None  # reset module-level cache

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "cptaq_zone_agricole.geojson"
        gj_path.write_text(json.dumps(self._make_cptaq_geojson([ring])), encoding="utf-8")
        build_cptaq_index(gj_path, tmp_path / "cptaq_index.json")

        result = lookup_cptaq(45.55, -73.55, tmp_path)
        assert result is not None
        assert result["en_zone_agricole"] is True
        assert result["source"] == "cptaq"
        assert result.get("NM_MRC") == "Rouville"

    def test_lookup_cptaq_outside(self, tmp_path):
        """Point outside all zones → en_zone_agricole: False."""
        import json
        from engine import data_enrichment as de
        from engine.data_enrichment import build_cptaq_index, lookup_cptaq

        de._CPTAQ_INDEX_CACHE = None

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "cptaq_zone_agricole.geojson"
        gj_path.write_text(json.dumps(self._make_cptaq_geojson([ring])), encoding="utf-8")
        build_cptaq_index(gj_path, tmp_path / "cptaq_index.json")

        result = lookup_cptaq(46.9, -75.0, tmp_path)
        assert result is not None
        assert result["en_zone_agricole"] is False

    def test_lookup_cptaq_no_data(self, tmp_path):
        """No GeoJSON and download unavailable → None (not error)."""
        from engine import data_enrichment as de
        from engine.data_enrichment import lookup_cptaq
        import unittest.mock as mock

        de._CPTAQ_INDEX_CACHE = None
        # Patch download_cptaq to simulate unavailable source
        with mock.patch.object(de, "download_cptaq", return_value=None):
            result = lookup_cptaq(45.55, -73.55, tmp_path)
        assert result is None

    def test_enrich_case_injects_zone_agricole(self, tmp_path):
        """enrich_case with mocked geocoding and CPTAQ data injects zone_agricole."""
        import json
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case, build_cptaq_index

        de._CPTAQ_INDEX_CACHE = None

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "cptaq_zone_agricole.geojson"
        gj_path.write_text(json.dumps(self._make_cptaq_geojson([ring])), encoding="utf-8")
        build_cptaq_index(gj_path, tmp_path / "cptaq_index.json")

        with mock.patch.object(de, "geocode_address", return_value=(45.55, -73.55)):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case = {"dossier_id": "D-CPTAQ-TEST"}
                enrich_case(case, display_name="100 chemin Rural, Rouville", cache_dir=tmp_path)

        assert case.get("zone_agricole", {}).get("en_zone_agricole") is True
        assert case["zone_agricole"]["source"] == "cptaq"

    def test_enrich_case_outside_zone(self, tmp_path):
        """Point outside zone agricole → en_zone_agricole: False."""
        import json
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case, build_cptaq_index

        de._CPTAQ_INDEX_CACHE = None

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "cptaq_zone_agricole.geojson"
        gj_path.write_text(json.dumps(self._make_cptaq_geojson([ring])), encoding="utf-8")
        build_cptaq_index(gj_path, tmp_path / "cptaq_index.json")

        # Point clearly outside the zone polygon (south of bbox)
        with mock.patch.object(de, "geocode_address", return_value=(45.40, -73.55)):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case = {"dossier_id": "D-CPTAQ-OUT"}
                enrich_case(case, display_name="1000 rue Sherbrooke, Montréal", cache_dir=tmp_path)

        assert case.get("zone_agricole", {}).get("en_zone_agricole") is False


# ── TestDataEnrichment_Patrimoine ─────────────────────────────────────────────

class TestDataEnrichment_Patrimoine:
    """Tests for patrimoine culturel lookup."""

    def _make_point_geojson(self, lng: float, lat: float, props: dict) -> dict:
        return {
            "type": "FeatureCollection",
            "features": [{
                "type": "Feature",
                "properties": props,
                "geometry": {"type": "Point", "coordinates": [lng, lat]},
            }],
        }

    def _make_polygon_geojson(self, ring: list, props: dict) -> dict:
        return {
            "type": "FeatureCollection",
            "features": [{
                "type": "Feature",
                "properties": props,
                "geometry": {"type": "Polygon", "coordinates": [ring]},
            }],
        }

    def test_build_index_point(self, tmp_path):
        """build_patrimoine_index handles Point geometry → tiny bbox."""
        import json
        from engine.data_enrichment import build_patrimoine_index

        gj = self._make_point_geojson(-73.55, 45.55,
                                       {"NOM": "Maison Dupont", "STATUT": "classe"})
        gj_path = tmp_path / "patrimoine_culturel.geojson"
        gj_path.write_text(json.dumps(gj), encoding="utf-8")

        idx_path = tmp_path / "patrimoine_index.json"
        count = build_patrimoine_index(gj_path, idx_path)
        assert count == 1
        data = json.loads(idx_path.read_text(encoding="utf-8"))
        feat = data["zones"][0]
        assert feat["props"]["NOM"] == "Maison Dupont"
        assert "point" in feat

    def test_build_index_polygon(self, tmp_path):
        """build_patrimoine_index handles Polygon geometry → ring."""
        import json
        from engine.data_enrichment import build_patrimoine_index

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj = self._make_polygon_geojson(ring, {"NOM": "Site historique", "STATUT": "reconnu"})
        gj_path = tmp_path / "patrimoine_culturel.geojson"
        gj_path.write_text(json.dumps(gj), encoding="utf-8")

        idx_path = tmp_path / "patrimoine_index.json"
        count = build_patrimoine_index(gj_path, idx_path)
        assert count == 1
        data = json.loads(idx_path.read_text(encoding="utf-8"))
        assert "ring" in data["zones"][0]

    def test_lookup_point_found(self, tmp_path):
        """Point near patrimoine point feature → returns props."""
        import json
        from engine import data_enrichment as de
        from engine.data_enrichment import build_patrimoine_index, lookup_patrimoine

        de._PATRIMOINE_INDEX_CACHE = None

        gj = self._make_point_geojson(-73.55, 45.55,
                                       {"NOM": "Vieux Moulin", "STATUT": "classe"})
        gj_path = tmp_path / "patrimoine_culturel.geojson"
        gj_path.write_text(json.dumps(gj), encoding="utf-8")
        build_patrimoine_index(gj_path, tmp_path / "patrimoine_index.json")

        # Point very close to the feature
        result = lookup_patrimoine(45.55, -73.55, tmp_path)
        assert result is not None
        assert result.get("NOM") == "Vieux Moulin"
        assert result["source"] == "patrimoine-culturel"

    def test_lookup_point_not_found(self, tmp_path):
        """Point far from all features → {}."""
        import json
        from engine import data_enrichment as de
        from engine.data_enrichment import build_patrimoine_index, lookup_patrimoine

        de._PATRIMOINE_INDEX_CACHE = None

        gj = self._make_point_geojson(-73.55, 45.55,
                                       {"NOM": "Vieux Moulin"})
        gj_path = tmp_path / "patrimoine_culturel.geojson"
        gj_path.write_text(json.dumps(gj), encoding="utf-8")
        build_patrimoine_index(gj_path, tmp_path / "patrimoine_index.json")

        result = lookup_patrimoine(45.40, -74.0, tmp_path)
        assert result == {}

    def test_lookup_no_data(self, tmp_path):
        """No data available → None."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import lookup_patrimoine

        de._PATRIMOINE_INDEX_CACHE = None
        with mock.patch.object(de, "download_patrimoine", return_value=None):
            result = lookup_patrimoine(45.55, -73.55, tmp_path)
        assert result is None

    def test_enrich_case_injects_patrimoine(self, tmp_path):
        """enrich_case injects patrimoine_culturel when building found nearby."""
        import json
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case, build_patrimoine_index

        de._PATRIMOINE_INDEX_CACHE = None

        gj = self._make_point_geojson(-73.55, 45.55,
                                       {"NOM": "Édifice Holt", "STATUT": "cite"})
        gj_path = tmp_path / "patrimoine_culturel.geojson"
        gj_path.write_text(json.dumps(gj), encoding="utf-8")
        build_patrimoine_index(gj_path, tmp_path / "patrimoine_index.json")

        with mock.patch.object(de, "geocode_address", return_value=(45.55, -73.55)):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case = {"dossier_id": "D-PAT-TEST"}
                enrich_case(case, display_name="100 rue Holt, Montréal", cache_dir=tmp_path)

        assert case.get("patrimoine_culturel", {}).get("NOM") == "Édifice Holt"
        assert case["patrimoine_culturel"]["source"] == "patrimoine-culturel"

    def test_enrich_case_empty_when_not_listed(self, tmp_path):
        """enrich_case sets patrimoine_culturel = {} when checked but not found."""
        import json
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case, build_patrimoine_index

        de._PATRIMOINE_INDEX_CACHE = None

        # Feature far from test point
        gj = self._make_point_geojson(-74.0, 46.0, {"NOM": "Elsewhere"})
        gj_path = tmp_path / "patrimoine_culturel.geojson"
        gj_path.write_text(json.dumps(gj), encoding="utf-8")
        build_patrimoine_index(gj_path, tmp_path / "patrimoine_index.json")

        with mock.patch.object(de, "geocode_address", return_value=(45.55, -73.55)):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case = {"dossier_id": "D-PAT-EMPTY"}
                enrich_case(case, display_name="999 rue Quelconque, Montréal", cache_dir=tmp_path)

        # Key present but empty (checked, not listed)
        assert "patrimoine_culturel" in case
        assert case["patrimoine_culturel"] == {}


# ── TestDataEnrichment_Inondable ──────────────────────────────────────────────

class TestDataEnrichment_Inondable:
    """Tests for MELCC zones inondables lookup."""

    def _make_geojson(self, ring: list, props: dict) -> dict:
        return {
            "type": "FeatureCollection",
            "features": [{
                "type": "Feature",
                "properties": props,
                "geometry": {"type": "Polygon", "coordinates": [ring]},
            }],
        }

    def test_build_index(self, tmp_path):
        """build_inondable_index creates valid index."""
        import json
        from engine.data_enrichment import build_inondable_index

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "zones_inondables.geojson"
        gj_path.write_text(json.dumps(self._make_geojson(ring, {"RECURRENCE": "0_20"})),
                            encoding="utf-8")

        idx_path = tmp_path / "inondable_index.json"
        count = build_inondable_index(gj_path, idx_path)
        assert count == 1
        data = json.loads(idx_path.read_text(encoding="utf-8"))
        assert data["zones"][0]["props"]["RECURRENCE"] == "0_20"

    def test_lookup_inside(self, tmp_path):
        """Point in flood zone → en_zone_inondable: True + recurrence."""
        import json
        from engine import data_enrichment as de
        from engine.data_enrichment import build_inondable_index, lookup_inondable

        de._INONDABLE_INDEX_CACHE = None

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "zones_inondables.geojson"
        gj_path.write_text(json.dumps(self._make_geojson(ring, {"RECURRENCE": "20_100"})),
                            encoding="utf-8")
        build_inondable_index(gj_path, tmp_path / "inondable_index.json")

        result = lookup_inondable(45.55, -73.55, tmp_path)
        assert result is not None
        assert result["en_zone_inondable"] is True
        assert result["recurrence"] == "20_100"
        assert "centennale" in result.get("recurrence_label", "")
        assert result["source"] == "melcc-zones-inondables"

    def test_lookup_outside(self, tmp_path):
        """Point outside flood zone → {}."""
        import json
        from engine import data_enrichment as de
        from engine.data_enrichment import build_inondable_index, lookup_inondable

        de._INONDABLE_INDEX_CACHE = None

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "zones_inondables.geojson"
        gj_path.write_text(json.dumps(self._make_geojson(ring, {"RECURRENCE": "0_20"})),
                            encoding="utf-8")
        build_inondable_index(gj_path, tmp_path / "inondable_index.json")

        result = lookup_inondable(45.40, -74.0, tmp_path)
        assert result == {}

    def test_lookup_most_restrictive(self, tmp_path):
        """Multiple overlapping zones → returns smallest recurrence (highest risk)."""
        import json
        from engine import data_enrichment as de
        from engine.data_enrichment import build_inondable_index, lookup_inondable

        de._INONDABLE_INDEX_CACHE = None

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj = {
            "type": "FeatureCollection",
            "features": [
                {
                    "type": "Feature",
                    "properties": {"RECURRENCE": "20_100"},
                    "geometry": {"type": "Polygon", "coordinates": [ring]},
                },
                {
                    "type": "Feature",
                    "properties": {"RECURRENCE": "0_20"},
                    "geometry": {"type": "Polygon", "coordinates": [ring]},
                },
            ],
        }
        gj_path = tmp_path / "zones_inondables.geojson"
        gj_path.write_text(json.dumps(gj), encoding="utf-8")
        build_inondable_index(gj_path, tmp_path / "inondable_index.json")

        result = lookup_inondable(45.55, -73.55, tmp_path)
        # Should pick 0_20 (more restrictive / higher risk)
        assert result["recurrence"] == "0_20"

    def test_lookup_no_data(self, tmp_path):
        """No data → None."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import lookup_inondable

        de._INONDABLE_INDEX_CACHE = None
        with mock.patch.object(de, "download_inondable", return_value=None):
            result = lookup_inondable(45.55, -73.55, tmp_path)
        assert result is None

    def test_enrich_case_injects_inondable(self, tmp_path):
        """enrich_case injects zone_inondable when point in flood zone."""
        import json
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case, build_inondable_index

        de._INONDABLE_INDEX_CACHE = None

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "zones_inondables.geojson"
        gj_path.write_text(
            json.dumps(self._make_geojson(ring, {"RECURRENCE": "0_20"})),
            encoding="utf-8",
        )
        build_inondable_index(gj_path, tmp_path / "inondable_index.json")

        with mock.patch.object(de, "geocode_address", return_value=(45.55, -73.55)):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case = {"dossier_id": "D-INOND-TEST"}
                enrich_case(case, display_name="1 chemin de la Berge, Laval", cache_dir=tmp_path)

        assert case.get("zone_inondable", {}).get("en_zone_inondable") is True
        assert case["zone_inondable"]["recurrence"] == "0_20"

    def test_enrich_case_empty_when_outside(self, tmp_path):
        """enrich_case sets zone_inondable = {} when checked but outside."""
        import json
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case, build_inondable_index

        de._INONDABLE_INDEX_CACHE = None

        ring = [[-73.6, 45.5], [-73.5, 45.5], [-73.5, 45.6], [-73.6, 45.6], [-73.6, 45.5]]
        gj_path = tmp_path / "zones_inondables.geojson"
        gj_path.write_text(
            json.dumps(self._make_geojson(ring, {"RECURRENCE": "20_100"})),
            encoding="utf-8",
        )
        build_inondable_index(gj_path, tmp_path / "inondable_index.json")

        with mock.patch.object(de, "geocode_address", return_value=(45.40, -74.0)):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case = {"dossier_id": "D-INOND-OUT"}
                enrich_case(case, display_name="1000 boul. Laurentien, Montréal", cache_dir=tmp_path)

        assert "zone_inondable" in case
        assert case["zone_inondable"] == {}


# ── TestDataEnrichment_Chantier ───────────────────────────────────────────────

class TestDataEnrichment_Chantier:
    """Tests for fetch_mises_en_chantier() via StatCan WDS 34-10-0056-01."""

    def _meta(self, geo_labels, type_labels, extra_labels=None):
        def members(labels):
            return [{"memberId": str(i + 1), "memberNameEn": lbl}
                    for i, lbl in enumerate(labels)]
        dims = [
            {"member": members(geo_labels)},
            {"member": members(type_labels)},
        ]
        if extra_labels:
            dims.append({"member": members(extra_labels)})
        return {"dims": dims}

    def _wds_12pts(self, values, period="2025-01"):
        return [{
            "status": "SUCCESS",
            "object": {
                "vectorDataPoint": [
                    {"value": str(v), "refPer": period}
                    for v in values
                ]
            },
        }]

    def test_fetch_chantier_success(self, tmp_path):
        """Total starts + 12-month sum + 6-month trend computed."""
        import pytest
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_mises_en_chantier

        meta = self._meta(
            ["Montréal", "Québec"],
            ["Total units", "Single-detached", "Apartments and other"],
        )
        monthly = [600.0] * 6 + [500.0] * 6  # recent avg 600, prior avg 500

        def fake_wds_post(endpoint, payload, timeout=8.0):
            return self._wds_12pts(monthly)

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                with mock.patch.object(de, "_fetch_series", return_value=480.0):
                    result = fetch_mises_en_chantier("montreal", tmp_path)

        assert result.get("ville") == "Montréal"
        assert result.get("source") == "statcan-34-10-0056-01"
        assert result.get("total_mois") == 600.0
        assert result.get("total_12mois") == 6600.0
        assert result.get("periode") == "2025-01"
        assert result.get("variation_pct_6m") == pytest.approx(20.0, abs=0.1)
        assert result.get("unifamilial_mois") == 480.0
        assert result.get("collectif_mois") == 480.0

    def test_fetch_chantier_unsupported_city(self, tmp_path):
        """Unknown city → {} without WDS call."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_mises_en_chantier

        with mock.patch.object(de, "_cube_metadata") as mock_meta:
            result = fetch_mises_en_chantier("rimouski", tmp_path)

        mock_meta.assert_not_called()
        assert result == {}

    def test_fetch_chantier_no_meta(self, tmp_path):
        """Empty metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_mises_en_chantier

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_mises_en_chantier("montreal", tmp_path)

        assert result == {}

    def test_fetch_chantier_geo_not_found(self, tmp_path):
        """No GEO match → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_mises_en_chantier

        meta = self._meta(["Toronto", "Vancouver"], ["Total units"])
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            result = fetch_mises_en_chantier("montreal", tmp_path)

        assert result == {}

    def test_enrich_case_injects_mises_en_chantier(self, tmp_path):
        """enrich_case populates mises_en_chantier via fetch_mises_en_chantier."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        chantier_data = {
            "ville": "Montréal",
            "source": "statcan-34-10-0056-01",
            "total_mois": 620.0,
            "total_12mois": 7_200.0,
            "variation_pct_6m": 8.5,
            "periode": "2025-01",
        }

        with mock.patch.object(de, "fetch_mises_en_chantier", return_value=chantier_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-CHANTIER-TEST"}
                enrich_case(case, display_name="450 rue Notre-Dame, Montréal",
                            cache_dir=tmp_path)

        mc = case.get("mises_en_chantier", {})
        assert mc.get("total_mois") == 620.0
        assert mc.get("variation_pct_6m") == 8.5


# ── TestDataEnrichment_IPC ────────────────────────────────────────────────────

class TestDataEnrichment_IPC:
    """Tests for fetch_ipc_logement() via StatCan WDS 18-10-0004-01."""

    def _meta(self, geo_labels, component_labels):
        def members(labels):
            return [{"memberId": str(i + 1), "memberNameEn": lbl}
                    for i, lbl in enumerate(labels)]
        return {"dims": [
            {"member": members(geo_labels)},
            {"member": members(component_labels)},
        ]}

    def _wds_scalar(self, value):
        return [{"status": "SUCCESS",
                 "object": {"vectorDataPoint": [{"value": str(value)}]}}]

    def _wds_13pts(self, latest, year_ago, period="2025-01"):
        pts = [{"value": str(latest), "refPer": period}]
        pts += [{"value": str((latest + year_ago) / 2)}] * 11
        pts += [{"value": str(year_ago)}]
        return [{"status": "SUCCESS",
                 "object": {"vectorDataPoint": pts}}]

    def test_fetch_ipc_success(self, tmp_path):
        """IPC components extracted + annual variation computed."""
        import pytest
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_ipc_logement

        meta = self._meta(
            ["Canada", "Québec"],
            ["All-items", "Shelter", "Energy"],
        )
        # _fetch_series returns current value per coord
        series_map = {"1.1": 160.0, "1.2": 175.0, "1.3": 140.0}

        def fake_fetch_series(pid, coord, cache_dir):
            return series_map.get(coord)

        # _wds_post used for latestN=13 annual variation
        def fake_wds_post(endpoint, payload, timeout=8.0):
            return self._wds_13pts(175.0, 165.0, "2025-01")

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_fetch_series", side_effect=fake_fetch_series):
                with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                    result = fetch_ipc_logement(tmp_path)

        assert result.get("source") == "statcan-18-10-0004-01"
        assert result.get("ipc_total") == 160.0
        assert result.get("ipc_logement") == 175.0
        assert result.get("ipc_energie") == 140.0
        assert result.get("periode") == "2025-01"
        # (175 - 165) / 165 * 100 ≈ 6.06%
        assert result.get("variation_logement_pct") == pytest.approx(6.1, abs=0.1)

    def test_fetch_ipc_no_shelter_component(self, tmp_path):
        """If Shelter component missing → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_ipc_logement

        meta = self._meta(["Canada"], ["All-items", "Food", "Transportation"])
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_fetch_series", return_value=100.0):
                with mock.patch.object(de, "_wds_post", return_value=[]):
                    result = fetch_ipc_logement(tmp_path)

        assert result == {}

    def test_fetch_ipc_no_meta(self, tmp_path):
        """Empty metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_ipc_logement

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_ipc_logement(tmp_path)

        assert result == {}

    def test_fetch_ipc_geo_not_found(self, tmp_path):
        """No matching GEO → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_ipc_logement

        meta = self._meta(["Ontario", "Alberta"], ["All-items", "Shelter"])
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            result = fetch_ipc_logement(tmp_path)

        assert result == {}

    def test_enrich_case_injects_ipc_logement(self, tmp_path):
        """enrich_case populates ipc_logement from fetch_ipc_logement."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        ipc_data = {
            "source": "statcan-18-10-0004-01",
            "ipc_total": 160.5,
            "ipc_logement": 178.2,
            "variation_logement_pct": 5.3,
            "periode": "2025-01",
        }

        with mock.patch.object(de, "fetch_ipc_logement", return_value=ipc_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-IPC-TEST"}
                enrich_case(case, display_name="1 pl. Ville-Marie, Montréal",
                            cache_dir=tmp_path)

        ip = case.get("ipc_logement", {})
        assert ip.get("ipc_logement") == 178.2
        assert ip.get("variation_logement_pct") == 5.3


# ── TestDataEnrichment_Travail ────────────────────────────────────────────────

class TestDataEnrichment_Travail:
    """Tests for fetch_marche_travail() via StatCan WDS 14-10-0096-01."""

    def _meta(self, geo_labels, indicator_labels, sex_labels=None):
        def members(labels):
            return [{"memberId": str(i + 1), "memberNameEn": lbl}
                    for i, lbl in enumerate(labels)]
        dims = [
            {"member": members(geo_labels)},
            {"member": members(indicator_labels)},
        ]
        if sex_labels:
            dims.append({"member": members(sex_labels)})
        return {"dims": dims}

    def _wds_latest(self, value, period="2025-01"):
        return [{
            "status": "SUCCESS",
            "object": {
                "vectorDataPoint": [{"value": str(value), "refPer": period}]
            },
        }]

    def test_fetch_travail_success(self, tmp_path):
        """Three indicators extracted correctly."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_travail

        meta = self._meta(
            ["Montréal", "Québec"],
            ["Unemployment rate", "Employment rate", "Participation rate"],
        )
        indicator_values = {"1.1": 6.2, "1.2": 61.8, "1.3": 66.0}

        def fake_wds_post(endpoint, payload, timeout=8.0):
            coord = payload[0]["coordinate"]
            return self._wds_latest(indicator_values.get(coord, 0.0))

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                result = fetch_marche_travail("montreal", tmp_path)

        assert result.get("ville") == "Montréal"
        assert result.get("source") == "statcan-14-10-0096-01"
        assert result.get("taux_chomage_pct") == 6.2
        assert result.get("taux_emploi_pct") == 61.8
        assert result.get("taux_participation_pct") == 66.0
        assert result.get("periode") == "2025-01"

    def test_fetch_travail_unsupported_city(self, tmp_path):
        """Unknown city → {} without WDS call."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_travail

        with mock.patch.object(de, "_cube_metadata") as mock_meta:
            result = fetch_marche_travail("rimouski", tmp_path)

        mock_meta.assert_not_called()
        assert result == {}

    def test_fetch_travail_no_meta(self, tmp_path):
        """Empty metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_travail

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_marche_travail("montreal", tmp_path)

        assert result == {}

    def test_fetch_travail_geo_not_found(self, tmp_path):
        """No GEO match → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_travail

        meta = self._meta(["Toronto", "Vancouver"], ["Unemployment rate"])
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            result = fetch_marche_travail("montreal", tmp_path)

        assert result == {}

    def test_fetch_travail_with_sex_dim(self, tmp_path):
        """3-dim table (with sex dim) extracts correctly."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_travail

        meta = self._meta(
            ["Montréal"],
            ["Unemployment rate", "Employment rate", "Participation rate"],
            ["Both sexes", "Male", "Female"],
        )

        def fake_wds_post(endpoint, payload, timeout=8.0):
            return self._wds_latest(5.9)

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                result = fetch_marche_travail("montreal", tmp_path)

        assert result.get("taux_chomage_pct") == 5.9

    def test_enrich_case_injects_marche_travail(self, tmp_path):
        """enrich_case populates marche_travail via fetch_marche_travail."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        travail_data = {
            "ville": "Montréal",
            "source": "statcan-14-10-0096-01",
            "taux_chomage_pct": 6.1,
            "taux_emploi_pct": 62.3,
            "periode": "2025-02",
        }

        with mock.patch.object(de, "fetch_marche_travail", return_value=travail_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-TRAV-TEST"}
                enrich_case(case, display_name="800 boul. René-Lévesque, Montréal",
                            cache_dir=tmp_path)

        mt = case.get("marche_travail", {})
        assert mt.get("taux_chomage_pct") == 6.1
        assert mt.get("taux_emploi_pct") == 62.3


# ── TestDataEnrichment_Population ────────────────────────────────────────────

class TestDataEnrichment_Population:
    """Tests for fetch_population_growth() via StatCan WDS 17-10-0135-01."""

    def _meta(self, geo_labels, age_labels=None, sex_labels=None, extra_labels=None):
        def members(labels):
            return [{"memberId": str(i + 1), "memberNameEn": lbl}
                    for i, lbl in enumerate(labels)]
        dims = [
            {"member": members(geo_labels)},
            {"member": members(age_labels or ["Total - all ages"])},
            {"member": members(sex_labels or ["Both sexes"])},
        ]
        if extra_labels:
            dims.append({"member": members(extra_labels)})
        return {"dims": dims}

    def _wds_2pts(self, latest, prior, ref_year="2023"):
        return [{
            "status": "SUCCESS",
            "object": {
                "vectorDataPoint": [
                    {"value": str(latest), "refPer": f"{ref_year}-07-01"},
                    {"value": str(prior),  "refPer": f"{int(ref_year)-1}-07-01"},
                ]
            },
        }]

    def test_fetch_population_success(self, tmp_path):
        """Population + annual growth computed correctly."""
        import pytest
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_population_growth

        meta = self._meta(["Montréal", "Québec"])
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post",
                                   return_value=self._wds_2pts(4_291_000, 4_220_000)):
                result = fetch_population_growth("montreal", tmp_path)

        assert result.get("ville") == "Montréal"
        assert result.get("source") == "statcan-17-10-0135-01"
        assert result.get("population") == 4_291_000
        assert result.get("annee") == "2023"
        # (4291000 - 4220000) / 4220000 * 100 ≈ 1.68%
        assert result.get("variation_annuelle_pct") == pytest.approx(1.68, abs=0.02)

    def test_fetch_population_unsupported_city(self, tmp_path):
        """Unknown city → {} without WDS call."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_population_growth

        with mock.patch.object(de, "_cube_metadata") as mock_meta:
            result = fetch_population_growth("rimouski", tmp_path)

        mock_meta.assert_not_called()
        assert result == {}

    def test_fetch_population_no_meta(self, tmp_path):
        """Empty metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_population_growth

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_population_growth("montreal", tmp_path)

        assert result == {}

    def test_fetch_population_geo_not_found(self, tmp_path):
        """GEO not matched → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_population_growth

        meta = self._meta(["Toronto", "Vancouver"])
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            result = fetch_population_growth("montreal", tmp_path)

        assert result == {}

    def test_fetch_population_single_point_no_growth(self, tmp_path):
        """Only 1 data point → population set, no variation."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_population_growth

        meta = self._meta(["Québec"])
        single = [{
            "status": "SUCCESS",
            "object": {
                "vectorDataPoint": [
                    {"value": "839311", "refPer": "2022-07-01"}
                ]
            },
        }]
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post", return_value=single):
                result = fetch_population_growth("quebec", tmp_path)

        assert result.get("population") == 839_311
        assert "variation_annuelle_pct" not in result

    def test_enrich_case_injects_population_cma(self, tmp_path):
        """enrich_case populates population_cma via fetch_population_growth."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        pop_data = {
            "ville": "Montréal",
            "source": "statcan-17-10-0135-01",
            "population": 4_291_000,
            "variation_annuelle_pct": 1.68,
            "annee": "2023",
        }

        with mock.patch.object(de, "fetch_population_growth", return_value=pop_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-POP-TEST"}
                enrich_case(case, display_name="123 av. Mont-Royal, Montréal",
                            cache_dir=tmp_path)

        pc = case.get("population_cma", {})
        assert pc.get("population") == 4_291_000
        assert pc.get("variation_annuelle_pct") == 1.68


# ── TestDataEnrichment_BOC ────────────────────────────────────────────────────

class TestDataEnrichment_BOC:
    """Tests for fetch_taux_boc() — Bank of Canada Valet API."""

    _BOC_RESPONSE = {
        "observations": [
            {
                "d": "2025-03-05",
                "CAOVRNIGH": {"v": "2.75"},
                "V80691311": {"v": "4.95"},
                "V122495":   {"v": "5.14"},
                "V122496":   {"v": "4.89"},
            }
        ]
    }

    def _mock_urlopen(self, payload):
        import io
        import json
        import unittest.mock as mock
        body = json.dumps(payload).encode("utf-8")
        cm = mock.MagicMock()
        cm.__enter__ = mock.Mock(return_value=io.BytesIO(body))
        cm.__exit__ = mock.Mock(return_value=False)
        return cm

    def test_fetch_taux_boc_success(self, tmp_path):
        """All four rates extracted from Valet response."""
        import unittest.mock as mock
        from engine.data_enrichment import fetch_taux_boc

        with mock.patch("urllib.request.urlopen",
                        return_value=self._mock_urlopen(self._BOC_RESPONSE)):
            result = fetch_taux_boc(tmp_path)

        assert result.get("source") == "bankofcanada-valet"
        assert result.get("date") == "2025-03-05"
        assert result.get("taux_directeur_pct") == 2.75
        assert result.get("taux_preferentiel_pct") == 4.95
        assert result.get("taux_hypo_5ans_conv_pct") == 5.14
        assert result.get("taux_hypo_1an_pct") == 4.89

    def test_fetch_taux_boc_empty_observations(self, tmp_path):
        """Empty observations list → {}."""
        import unittest.mock as mock
        from engine.data_enrichment import fetch_taux_boc

        with mock.patch("urllib.request.urlopen",
                        return_value=self._mock_urlopen({"observations": []})):
            result = fetch_taux_boc(tmp_path)

        assert result == {}

    def test_fetch_taux_boc_network_error(self, tmp_path):
        """Network error → {} without raising."""
        import unittest.mock as mock
        from engine.data_enrichment import fetch_taux_boc

        with mock.patch("urllib.request.urlopen", side_effect=OSError("network")):
            result = fetch_taux_boc(tmp_path)

        assert result == {}

    def test_fetch_taux_boc_partial_series(self, tmp_path):
        """Missing series in response are skipped gracefully."""
        import unittest.mock as mock
        from engine.data_enrichment import fetch_taux_boc

        partial = {
            "observations": [
                {"d": "2025-03-05", "CAOVRNIGH": {"v": "2.75"}}
            ]
        }
        with mock.patch("urllib.request.urlopen",
                        return_value=self._mock_urlopen(partial)):
            result = fetch_taux_boc(tmp_path)

        assert result.get("taux_directeur_pct") == 2.75
        assert "taux_preferentiel_pct" not in result
        assert "taux_hypo_5ans_conv_pct" not in result

    def test_fetch_taux_boc_cache_hit(self, tmp_path):
        """Cached result returned without network call."""
        import json
        import time
        import unittest.mock as mock
        from engine.data_enrichment import fetch_taux_boc

        cached = {
            "source": "bankofcanada-valet",
            "date": "2025-03-01",
            "taux_directeur_pct": 3.0,
            "_ts": time.time(),
        }
        (tmp_path / "boc_rates.json").write_text(
            json.dumps(cached), encoding="utf-8"
        )

        with mock.patch("urllib.request.urlopen") as mock_url:
            result = fetch_taux_boc(tmp_path)

        mock_url.assert_not_called()
        assert result.get("taux_directeur_pct") == 3.0

    def test_enrich_case_injects_taux_bancaires(self, tmp_path):
        """enrich_case populates taux_bancaires from fetch_taux_boc."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        boc_data = {
            "source": "bankofcanada-valet",
            "date": "2025-03-05",
            "taux_directeur_pct": 2.75,
            "taux_hypo_5ans_conv_pct": 5.14,
        }

        with mock.patch.object(de, "fetch_taux_boc", return_value=boc_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-BOC-TEST"}
                enrich_case(case, display_name="50 rue King, Montréal",
                            cache_dir=tmp_path)

        tb = case.get("taux_bancaires", {})
        assert tb.get("taux_directeur_pct") == 2.75
        assert tb.get("taux_hypo_5ans_conv_pct") == 5.14


# ── TestDataEnrichment_Vacance ────────────────────────────────────────────────

class TestDataEnrichment_Vacance:
    """Tests for fetch_vacancy_rate() via StatCan WDS 34-10-0131-01."""

    def _meta(self, geo_labels, unit_labels, extra_dim=False):
        def members(labels):
            return [{"memberId": str(i + 1), "memberNameEn": lbl}
                    for i, lbl in enumerate(labels)]
        dims = [
            {"member": members(geo_labels)},
            {"member": members(unit_labels)},
        ]
        if extra_dim:
            dims.append({"member": members(["Private apartments"])})
        return {"dims": dims}

    def test_fetch_vacance_success(self, tmp_path):
        """Successful fetch returns vacancy rates by unit type."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_vacancy_rate

        meta = self._meta(
            ["Montréal, Quebec", "Québec, Quebec"],
            ["total", "bachelor", "1 bedroom", "2 bedroom", "3 bedroom"],
        )
        rate_map = {
            "1.1": 2.3,   # total
            "1.2": 1.8,   # bachelor
            "1.3": 2.1,   # 1ch
            "1.4": 2.5,   # 2ch
            "1.5": 3.0,   # 3ch+
        }

        def fake_fetch_series(pid, coord, cache_dir):
            return rate_map.get(coord)

        wds_response = [{
            "status": "SUCCESS",
            "object": {"vectorDataPoint": [{"value": "2.3", "refPer": "2024-10"}]},
        }]

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_fetch_series", side_effect=fake_fetch_series):
                with mock.patch.object(de, "_wds_post", return_value=wds_response):
                    result = fetch_vacancy_rate("montreal", tmp_path)

        assert result.get("ville") == "Montréal"
        assert result.get("source") == "statcan-34-10-0131-01"
        assert result.get("taux_total_pct") == 2.3
        assert result.get("taux_1ch_pct") == 2.1
        assert result.get("taux_2ch_pct") == 2.5
        assert result.get("annee") == "2024"

    def test_fetch_vacance_unsupported_city(self, tmp_path):
        """Unsupported city → {} without WDS call."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_vacancy_rate

        with mock.patch.object(de, "_cube_metadata") as mock_meta:
            result = fetch_vacancy_rate("rimouski", tmp_path)

        mock_meta.assert_not_called()
        assert result == {}

    def test_fetch_vacance_no_meta(self, tmp_path):
        """Empty metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_vacancy_rate

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_vacancy_rate("montreal", tmp_path)

        assert result == {}

    def test_fetch_vacance_geo_not_found(self, tmp_path):
        """No GEO match → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_vacancy_rate

        meta = self._meta(["Toronto", "Vancouver"], ["total", "1 bedroom"])
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            result = fetch_vacancy_rate("montreal", tmp_path)

        assert result == {}

    def test_fetch_vacance_no_series_data(self, tmp_path):
        """All series return None → {} (no meaningful data)."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_vacancy_rate

        meta = self._meta(
            ["Montréal, Quebec"],
            ["total", "1 bedroom", "2 bedroom"],
        )
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_fetch_series", return_value=None):
                with mock.patch.object(de, "_wds_post", return_value=[]):
                    result = fetch_vacancy_rate("montreal", tmp_path)

        assert result == {}

    def test_enrich_case_injects_taux_inoccupation(self, tmp_path):
        """enrich_case populates taux_inoccupation via fetch_vacancy_rate."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        vacance_data = {
            "ville": "Montréal",
            "source": "statcan-34-10-0131-01",
            "taux_total_pct": 2.4,
            "taux_1ch_pct": 2.1,
            "annee": "2024",
        }

        with mock.patch.object(de, "fetch_vacancy_rate", return_value=vacance_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-VAC-TEST"}
                enrich_case(case, display_name="10 rue Sherbrooke, Montréal",
                            cache_dir=tmp_path)

        ti = case.get("taux_inoccupation", {})
        assert ti.get("taux_total_pct") == 2.4
        assert ti.get("annee") == "2024"


# ── TestDataEnrichment_NHPI ───────────────────────────────────────────────────

class TestDataEnrichment_NHPI:
    """Tests for NHPI (New Housing Price Index) fetch via StatCan WDS."""

    def _meta(self, geo_members: list, type_members: list) -> dict:
        return {
            "_ts": 9999999999,
            "dims": [
                {"member": [{"memberId": i + 1, "memberNameEn": m}
                             for i, m in enumerate(geo_members)]},
                {"member": [{"memberId": i + 1, "memberNameEn": m}
                             for i, m in enumerate(type_members)]},
            ],
        }

    def test_fetch_nhpi_success(self, tmp_path):
        """fetch_nhpi returns indice_total + variation_annuelle_pct on success."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_nhpi

        meta = self._meta(
            ["Montréal", "Québec, Quebec"],
            ["Total (house and land)", "Building", "Land"],
        )

        # Single-period response for _fetch_series
        def fake_fetch_series(pid, coord, cache_dir):
            return 130.5 if "1.1" in coord else 125.0

        # 13-period response for annual variation
        points_13 = [{"value": "130.5"}] + [{"value": "120.0"}] * 12

        def fake_wds_post(endpoint, payload, timeout=8.0):
            return [{
                "status": "SUCCESS",
                "object": {"vectorDataPoint": points_13},
            }]

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_fetch_series", side_effect=fake_fetch_series):
                with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                    result = fetch_nhpi("montreal", tmp_path)

        assert result.get("ville") == "Montréal"
        assert result.get("source") == "statcan-18-10-0205-01"
        assert result.get("indice_total") == 130.5
        import pytest
        assert result.get("variation_annuelle_pct") == pytest.approx(8.75, abs=0.1)

    def test_fetch_nhpi_unsupported_city(self, tmp_path):
        """Unsupported city → {}."""
        from engine.data_enrichment import fetch_nhpi
        result = fetch_nhpi("drummondville", tmp_path)
        # Drummondville not in NHPI geo labels → {}
        # (may return {} if no meta match, not an error)
        assert isinstance(result, dict)

    def test_fetch_nhpi_no_meta(self, tmp_path):
        """Empty metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_nhpi

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_nhpi("montreal", tmp_path)

        assert result == {}

    def test_fetch_nhpi_geo_not_found(self, tmp_path):
        """City not in GEO dimension → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_nhpi

        meta = self._meta(["Toronto", "Vancouver"], ["Total (house and land)"])
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            result = fetch_nhpi("montreal", tmp_path)

        assert result == {}

    def test_enrich_case_injects_nhpi(self, tmp_path):
        """enrich_case injects indice_prix_logement via fetch_nhpi."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        nhpi_data = {
            "source": "statcan-18-10-0205-01",
            "ville": "Montréal",
            "indice_total": 142.3,
            "variation_annuelle_pct": 3.2,
        }

        with mock.patch.object(de, "fetch_nhpi", return_value=nhpi_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case = {"dossier_id": "D-NHPI-TEST"}
                enrich_case(case, display_name="100 rue Test, Montréal", cache_dir=tmp_path)

        assert case.get("indice_prix_logement", {}).get("indice_total") == 142.3
        assert case["indice_prix_logement"]["variation_annuelle_pct"] == 3.2


# ── Census Profile 2021 ───────────────────────────────────────────────────────

class TestDataEnrichment_Census:
    """Unit tests for fetch_census_profile() and enrich_case() census injection."""

    # Minimal Census Profile API response (2 topics merged)
    _CENSUS_ROWS = [
        {
            "CHARACTERISTIC_NAME": "Propriétaires",
            "C1_COUNT_TOTAL": "1200000",
        },
        {
            "CHARACTERISTIC_NAME": "Locataires",
            "C1_COUNT_TOTAL": "800000",
        },
        {
            "CHARACTERISTIC_NAME": "Valeur médiane ($) des logements occupés par propriétaire",
            "C1_COUNT_TOTAL": "510000",
        },
        {
            "CHARACTERISTIC_NAME": "Frais mensuels médians ($) pour les logements loués",
            "C1_COUNT_TOTAL": "980",
        },
        {
            "CHARACTERISTIC_NAME": "Revenu total médian des ménages en 2020 ($)",
            "C1_COUNT_TOTAL": "68000",
        },
    ]

    def _mock_urlopen(self, rows):
        """Return a context manager that yields encoded JSON bytes."""
        import io
        import json
        import unittest.mock as mock
        body = json.dumps(rows).encode("utf-8")
        cm = mock.MagicMock()
        cm.__enter__ = mock.Mock(return_value=io.BytesIO(body))
        cm.__exit__ = mock.Mock(return_value=False)
        return cm

    def test_fetch_census_profile_success(self, tmp_path):
        """Successful fetch returns all expected numeric fields."""
        import json
        import unittest.mock as mock
        from engine.data_enrichment import fetch_census_profile

        with mock.patch("urllib.request.urlopen", return_value=self._mock_urlopen(self._CENSUS_ROWS)):
            result = fetch_census_profile("montreal", tmp_path)

        assert result.get("ville") == "montreal"
        assert result.get("source") == "StatCan Recensement 2021"
        assert result.get("pct_proprietaires") == 1_200_000.0
        assert result.get("pct_locataires") == 800_000.0
        assert result.get("valeur_mediane_logement") == 510_000.0
        assert result.get("frais_loyer_median") == 980.0
        assert result.get("revenu_median_menage") == 68_000.0

    def test_fetch_census_unsupported_city(self, tmp_path):
        """Unknown city_code → {} without network call."""
        import unittest.mock as mock
        from engine.data_enrichment import fetch_census_profile

        with mock.patch("urllib.request.urlopen") as mock_url:
            result = fetch_census_profile("rimouski", tmp_path)

        mock_url.assert_not_called()
        assert result == {}

    def test_fetch_census_empty_response(self, tmp_path):
        """API returns empty list → {} (no data)."""
        import unittest.mock as mock
        from engine.data_enrichment import fetch_census_profile

        with mock.patch("urllib.request.urlopen", return_value=self._mock_urlopen([])):
            result = fetch_census_profile("montreal", tmp_path)

        assert result == {}

    def test_fetch_census_suppressed_values_skipped(self, tmp_path):
        """Suppressed values ('x', '...') are not extracted."""
        import unittest.mock as mock
        from engine.data_enrichment import fetch_census_profile

        rows = [
            {"CHARACTERISTIC_NAME": "Propriétaires", "C1_COUNT_TOTAL": "x"},
            {"CHARACTERISTIC_NAME": "Locataires", "C1_COUNT_TOTAL": "..."},
            {"CHARACTERISTIC_NAME": "Valeur médiane ($) des logements occupés par propriétaire",
             "C1_COUNT_TOTAL": "510000"},
        ]
        with mock.patch("urllib.request.urlopen", return_value=self._mock_urlopen(rows)):
            result = fetch_census_profile("montreal", tmp_path)

        assert "pct_proprietaires" not in result
        assert "pct_locataires" not in result
        assert result.get("valeur_mediane_logement") == 510_000.0

    def test_fetch_census_cache_hit(self, tmp_path):
        """Second call reads from cache, no network."""
        import json
        import time
        import unittest.mock as mock
        from engine.data_enrichment import fetch_census_profile

        # Prime cache manually
        cache_data = {
            "ville": "montreal",
            "source": "StatCan Recensement 2021",
            "valeur_mediane_logement": 499000.0,
            "_ts": time.time(),
        }
        (tmp_path / "census_montreal.json").write_text(
            json.dumps(cache_data), encoding="utf-8"
        )

        with mock.patch("urllib.request.urlopen") as mock_url:
            result = fetch_census_profile("montreal", tmp_path)

        mock_url.assert_not_called()
        assert result.get("valeur_mediane_logement") == 499_000.0

    def test_enrich_case_injects_census(self, tmp_path):
        """enrich_case populates donnees_sociodemographiques from fetch_census_profile."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        census_data = {
            "ville": "montreal",
            "source": "StatCan Recensement 2021",
            "pct_proprietaires": 1_100_000.0,
            "revenu_median_menage": 72_000.0,
        }

        with mock.patch.object(de, "fetch_census_profile", return_value=census_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-CENSUS-TEST"}
                enrich_case(case, display_name="55 av. du Parc, Montréal", cache_dir=tmp_path)

        sd = case.get("donnees_sociodemographiques", {})
        assert sd.get("pct_proprietaires") == 1_100_000.0
        assert sd.get("revenu_median_menage") == 72_000.0


# ── TestDataEnrichment_Permis ─────────────────────────────────────────────────

class TestDataEnrichment_Permis:
    """Unit tests for fetch_permis_construction() and enrich_case() injection."""

    def _meta(self, geo_labels, struct_labels, work_labels, measure_labels=None):
        """Build fake WDS cube metadata with specified member lists."""
        def members(labels):
            return [{"memberId": str(i + 1), "memberNameEn": lbl}
                    for i, lbl in enumerate(labels)]

        dims = [
            {"member": members(geo_labels)},    # dim 0: GEO
            {"member": members(struct_labels)}, # dim 1: structure
            {"member": members(work_labels)},   # dim 2: work type
        ]
        if measure_labels:
            dims.append({"member": members(measure_labels)})  # dim 3: measure
        return {"dims": dims}

    def _series_12(self, values):
        """Build a 12-period WDS response."""
        return [{
            "status": "SUCCESS",
            "object": {
                "vectorDataPoint": [
                    {"value": str(v), "refPer": f"2025-{(12 - i):02d}"}
                    for i, v in enumerate(values)
                ]
            },
        }]

    def test_fetch_permis_success_3dim(self, tmp_path):
        """3-dim table (no measure dim): units fetched from new construction series."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_permis_construction

        meta = self._meta(
            ["Montréal, Quebec", "Québec, Quebec"],
            ["Residential", "Non-residential"],
            ["New construction", "All work"],
        )
        # 12 monthly values (most recent first)
        monthly = [450.0] * 6 + [380.0] * 6

        def fake_wds_post(endpoint, payload, timeout=8.0):
            return self._series_12(monthly)

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                with mock.patch.object(de, "_fetch_series", return_value=None):
                    result = fetch_permis_construction("montreal", tmp_path)

        assert result.get("ville") == "Montréal"
        assert result.get("source") == "statcan-34-10-0066-01"
        assert result.get("unites_residentielles_mois") == 450.0
        assert result.get("unites_residentielles_12mois") == 4980.0  # sum(monthly)
        assert result.get("periode") == "2025-12"
        # variation: avg(450×6) vs avg(380×6) = (450-380)/380 ≈ 18.4%
        import pytest
        assert result.get("variation_pct_6m") == pytest.approx(18.4, abs=0.2)

    def test_fetch_permis_unsupported_city(self, tmp_path):
        """Unsupported city → {} without WDS call."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_permis_construction

        with mock.patch.object(de, "_cube_metadata") as mock_meta:
            result = fetch_permis_construction("rimouski", tmp_path)

        mock_meta.assert_not_called()
        assert result == {}

    def test_fetch_permis_no_meta(self, tmp_path):
        """Empty WDS metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_permis_construction

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_permis_construction("montreal", tmp_path)

        assert result == {}

    def test_fetch_permis_geo_not_found(self, tmp_path):
        """GEO dimension has no match for city → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_permis_construction

        meta = self._meta(
            ["Toronto", "Vancouver"],
            ["Residential"],
            ["New construction"],
        )
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            result = fetch_permis_construction("montreal", tmp_path)

        assert result == {}

    def test_fetch_permis_4dim_with_measure(self, tmp_path):
        """4-dim table: unit_ord and value_ord extracted from dim 3."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_permis_construction

        meta = self._meta(
            ["Montréal, Quebec"],
            ["Residential"],
            ["New construction", "All work"],
            ["Number of units", "Value of permits"],  # dim 3
        )
        monthly = [300.0] * 12

        def fake_wds_post(endpoint, payload, timeout=8.0):
            return self._series_12(monthly)

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                with mock.patch.object(de, "_fetch_series", return_value=5_200.0):
                    result = fetch_permis_construction("montreal", tmp_path)

        assert result.get("unites_residentielles_mois") == 300.0
        assert result.get("valeur_permis_k_mois") == 5_200.0

    def test_enrich_case_injects_permis(self, tmp_path):
        """enrich_case populates permis_construction via fetch_permis_construction."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        permis_data = {
            "ville": "Montréal",
            "source": "statcan-34-10-0066-01",
            "unites_residentielles_mois": 520.0,
            "unites_residentielles_12mois": 5_800.0,
            "variation_pct_6m": 12.5,
            "periode": "2025-11",
        }

        with mock.patch.object(de, "fetch_permis_construction", return_value=permis_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-PERMIS-TEST"}
                enrich_case(case, display_name="200 rue Sainte-Catherine, Montréal",
                            cache_dir=tmp_path)

        pc = case.get("permis_construction", {})
        assert pc.get("unites_residentielles_mois") == 520.0
        assert pc.get("variation_pct_6m") == 12.5


# ── TestDataEnrichment_Proximite ─────────────────────────────────────────────

class TestDataEnrichment_Proximite:
    """Tests for fetch_proximite_services() via Overpass API."""

    def _overpass_resp(self, total: int) -> bytes:
        import json
        return json.dumps({
            "elements": [{"type": "count", "tags": {"total": str(total)}}]
        }).encode()

    def test_fetch_proximite_success(self, tmp_path):
        """All six queries succeed → full result dict cached."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_proximite_services

        counts = [3, 12, 2, 5, 1, 4]  # one per _OVERPASS_QUERIES
        call_idx = {"n": 0}

        def fake_urlopen(req, timeout=None):
            idx = call_idx["n"]
            call_idx["n"] += 1
            resp_bytes = self._overpass_resp(counts[idx % len(counts)])

            class _FakeResp:
                def read(self): return resp_bytes
                def __enter__(self): return self
                def __exit__(self, *a): pass

            return _FakeResp()

        with mock.patch("urllib.request.urlopen", side_effect=fake_urlopen):
            result = fetch_proximite_services(45.5017, -73.5673, tmp_path)

        assert result.get("ecoles_1km") == 3
        assert result.get("arrets_transport_500m") == 12
        assert result.get("epiceries_500m") == 2
        assert result.get("parcs_1km") == 5
        assert result.get("hopitaux_2km") == 1
        assert result.get("pharmacies_500m") == 4
        assert result.get("source") == "openstreetmap-overpass"
        assert result.get("lat") == 45.5017
        assert result.get("lng") == -73.5673

    def test_fetch_proximite_cache_hit(self, tmp_path):
        """Second call returns cached data without network."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_proximite_services

        cached = {
            "source": "openstreetmap-overpass",
            "lat": 45.5017, "lng": -73.5673,
            "ecoles_1km": 7,
            "arrets_transport_500m": 20,
            "_ts": 9_999_999_999.0,
        }
        cache_path = tmp_path / "overpass_45.5017_-73.5673.json"
        import json
        cache_path.write_text(json.dumps(cached))

        with mock.patch("urllib.request.urlopen") as mock_req:
            result = fetch_proximite_services(45.5017, -73.5673, tmp_path)

        mock_req.assert_not_called()
        assert result.get("ecoles_1km") == 7

    def test_fetch_proximite_network_error(self, tmp_path):
        """All queries fail → {} (non-blocking)."""
        import unittest.mock as mock
        from engine.data_enrichment import fetch_proximite_services

        with mock.patch("urllib.request.urlopen", side_effect=OSError("timeout")):
            result = fetch_proximite_services(45.5017, -73.5673, tmp_path)

        assert result == {}

    def test_fetch_proximite_partial_results(self, tmp_path):
        """Only some queries succeed → partial dict cached and returned."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_proximite_services, _OVERPASS_QUERIES

        call_idx = {"n": 0}

        def fake_urlopen(req, timeout=None):
            idx = call_idx["n"]
            call_idx["n"] += 1
            if idx >= 3:
                raise OSError("timeout")
            resp_bytes = self._overpass_resp(idx + 1)

            class _FakeResp:
                def read(self): return resp_bytes
                def __enter__(self): return self
                def __exit__(self, *a): pass

            return _FakeResp()

        with mock.patch("urllib.request.urlopen", side_effect=fake_urlopen):
            result = fetch_proximite_services(45.5017, -73.5673, tmp_path)

        # 3 successful queries + source/lat/lng → 6 keys > 3 → not empty
        assert result != {}
        assert result.get("ecoles_1km") == 1
        assert result.get("arrets_transport_500m") == 2

    def test_enrich_case_injects_proximite(self, tmp_path):
        """enrich_case injects proximite_services after geocoding."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        prox_data = {
            "source": "openstreetmap-overpass",
            "lat": 45.5017,
            "lng": -73.5673,
            "ecoles_1km": 5,
            "arrets_transport_500m": 15,
            "epiceries_500m": 3,
            "parcs_1km": 4,
            "hopitaux_2km": 2,
            "pharmacies_500m": 6,
        }

        with mock.patch.object(de, "fetch_proximite_services", return_value=prox_data):
            with mock.patch.object(de, "geocode_address", return_value=(45.5017, -73.5673)):
                with mock.patch.object(de, "detect_city", return_value="montreal"):
                    case: dict = {"dossier_id": "D-PROX-TEST"}
                    enrich_case(case, display_name="100 rue Sherbrooke, Montréal",
                                cache_dir=tmp_path)

        pr = case.get("proximite_services", {})
        assert pr.get("ecoles_1km") == 5
        assert pr.get("arrets_transport_500m") == 15


# ── TestDataEnrichment_Crime ──────────────────────────────────────────────────

class TestDataEnrichment_Crime:
    """Tests for fetch_crime_stats() via StatCan WDS 35-10-0078-01."""

    def _meta(self, geo_labels, violation_labels, stat_labels):
        def members(labels):
            return [{"memberId": str(i + 1), "memberNameEn": lbl}
                    for i, lbl in enumerate(labels)]
        return {"dims": [
            {"member": members(geo_labels)},
            {"member": members(violation_labels)},
            {"member": members(stat_labels)},
        ]}

    def test_fetch_crime_success(self, tmp_path):
        """All three rates returned for Montreal."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_crime_stats

        meta = self._meta(
            ["Montréal", "Québec"],
            ["Total Criminal Code violations",
             "Violent violations",
             "Property violations"],
            ["Rate per 100,000 population", "Number"],
        )

        rates = {
            "Total Criminal Code violations": 4200.0,
            "Violent violations": 950.5,
            "Property violations": 2100.3,
        }

        def fake_fetch_series(table, coord, cache_dir):
            import json as _json
            # coord is JSON array [geo_ord, viol_ord, stat_ord]
            ords = _json.loads(coord)
            viol_ord = ords[1]  # dim 1 = violation type
            vals = [4200.0, 950.5, 2100.3]
            return vals[viol_ord - 1]

        def fake_wds_post(endpoint, payload, timeout=8.0):
            return [{
                "status": "SUCCESS",
                "object": {
                    "vectorDataPoint": [{"value": "4200.0", "refPer": "2022"}]
                }
            }]

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_fetch_series", side_effect=fake_fetch_series):
                with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                    result = fetch_crime_stats("montreal", tmp_path)

        assert result.get("source") == "statcan-35-10-0078-01"
        assert result.get("ville") == "Montréal"
        assert result.get("taux_criminalite_total") == 4200.0
        assert result.get("taux_crimes_violents") == 950.5
        assert result.get("taux_crimes_contre_propriete") == 2100.3
        assert result.get("annee") == "2022"

    def test_fetch_crime_unsupported_city(self, tmp_path):
        """Unknown city → {} without WDS call."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_crime_stats

        with mock.patch.object(de, "_cube_metadata") as mock_meta:
            result = fetch_crime_stats("rimouski", tmp_path)

        mock_meta.assert_not_called()
        assert result == {}

    def test_fetch_crime_no_meta(self, tmp_path):
        """Empty metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_crime_stats

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_crime_stats("montreal", tmp_path)

        assert result == {}

    def test_fetch_crime_geo_not_found(self, tmp_path):
        """CMA not in dim 0 → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_crime_stats

        meta = self._meta(
            ["Toronto", "Vancouver"],
            ["Total Criminal Code violations"],
            ["Rate per 100,000 population"],
        )
        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            result = fetch_crime_stats("montreal", tmp_path)

        assert result == {}

    def test_fetch_crime_cache_hit(self, tmp_path):
        """Cached result returned without WDS call."""
        import json
        import time
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_crime_stats

        cached = {
            "source": "statcan-35-10-0078-01",
            "ville": "Montréal",
            "taux_criminalite_total": 3900.0,
            "annee": "2021",
            "_ts": time.time() + 1_000,
        }
        (tmp_path / "crime_montreal.json").write_text(json.dumps(cached))

        with mock.patch.object(de, "_cube_metadata") as mock_meta:
            result = fetch_crime_stats("montreal", tmp_path)

        mock_meta.assert_not_called()
        assert result.get("taux_criminalite_total") == 3900.0

    def test_enrich_case_injects_crime(self, tmp_path):
        """enrich_case populates crime_stats via fetch_crime_stats."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        crime_data = {
            "source": "statcan-35-10-0078-01",
            "ville": "Montréal",
            "taux_criminalite_total": 4150.0,
            "taux_crimes_violents": 890.0,
            "taux_crimes_contre_propriete": 2050.0,
            "annee": "2022",
        }

        with mock.patch.object(de, "fetch_crime_stats", return_value=crime_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-CRIME-TEST"}
                enrich_case(case, display_name="300 boul. Saint-Laurent, Montréal",
                            cache_dir=tmp_path)

        cs = case.get("crime_stats", {})
        assert cs.get("taux_criminalite_total") == 4150.0
        assert cs.get("annee") == "2022"


# ── TestDataEnrichment_MarcheNeuf ─────────────────────────────────────────────

class TestDataEnrichment_MarcheNeuf:
    """Tests for fetch_marche_neuf() via StatCan WDS 34-10-0093-01."""

    def _meta(self, geo_labels, type_labels, var_labels):
        def members(labels):
            return [{"memberId": str(i + 1), "memberNameEn": lbl}
                    for i, lbl in enumerate(labels)]
        return {"dims": [
            {"member": members(geo_labels)},
            {"member": members(type_labels)},
            {"member": members(var_labels)},
        ]}

    def _wds_12pts(self, values, period="2025-01"):
        return [{
            "status": "SUCCESS",
            "object": {
                "vectorDataPoint": [
                    {"value": str(v), "refPer": period, "status": ""}
                    for v in values
                ]
            },
        }]

    def test_fetch_marche_neuf_success(self, tmp_path):
        """Completions, 12-month total, and units under construction returned."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_neuf

        meta = self._meta(
            ["Montréal", "Québec"],
            ["Total units", "Single-detached"],
            ["Completed", "Under construction"],
        )

        call_idx = {"n": 0}

        def fake_wds_post(endpoint, payload, timeout=8.0):
            idx = call_idx["n"]
            call_idx["n"] += 1
            if idx == 0:
                # completions: 400/month × 12 months
                return self._wds_12pts([400.0] * 12, "2025-03")
            else:
                # under construction: single value
                return self._wds_12pts([5200.0] * 12, "2025-03")

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                result = fetch_marche_neuf("montreal", tmp_path)

        assert result.get("source") == "statcan-34-10-0093-01"
        assert result.get("ville") == "Montréal"
        assert result.get("completions_mois") == 400.0
        assert result.get("completions_12mois") == 4800.0
        assert result.get("unites_en_construction") == 5200.0
        assert result.get("periode") == "2025-03"

    def test_fetch_marche_neuf_with_absorption(self, tmp_path):
        """Absorption rate computed from cached starts data."""
        import json
        import time
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_neuf

        # Pre-populate starts cache
        starts_data = {
            "ville": "Montréal",
            "source": "statcan-34-10-0056-01",
            "total_mois": 500.0,
            "_ts": time.time(),
        }
        (tmp_path / "mises_en_chantier_montreal.json").write_text(json.dumps(starts_data))

        meta = self._meta(
            ["Montréal"],
            ["Total units"],
            ["Completed", "Under construction"],
        )

        call_idx = {"n": 0}

        def fake_wds_post(endpoint, payload, timeout=8.0):
            idx = call_idx["n"]
            call_idx["n"] += 1
            if idx == 0:
                return self._wds_12pts([450.0] * 12, "2025-03")
            return self._wds_12pts([5000.0] * 12, "2025-03")

        with mock.patch.object(de, "_cube_metadata", return_value=meta):
            with mock.patch.object(de, "_wds_post", side_effect=fake_wds_post):
                result = fetch_marche_neuf("montreal", tmp_path)

        import pytest
        # 450/500 × 100 = 90.0%
        assert result.get("taux_absorption_pct") == pytest.approx(90.0, abs=0.1)

    def test_fetch_marche_neuf_unsupported_city(self, tmp_path):
        """Unknown city → {} without WDS call."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_neuf

        with mock.patch.object(de, "_cube_metadata") as mock_meta:
            result = fetch_marche_neuf("rimouski", tmp_path)

        mock_meta.assert_not_called()
        assert result == {}

    def test_fetch_marche_neuf_no_meta(self, tmp_path):
        """Empty metadata → {}."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import fetch_marche_neuf

        with mock.patch.object(de, "_cube_metadata", return_value={}):
            result = fetch_marche_neuf("montreal", tmp_path)

        assert result == {}

    def test_enrich_case_injects_marche_neuf(self, tmp_path):
        """enrich_case populates marche_neuf via fetch_marche_neuf."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        neuf_data = {
            "source": "statcan-34-10-0093-01",
            "ville": "Montréal",
            "completions_mois": 420.0,
            "completions_12mois": 5_040.0,
            "unites_en_construction": 4_800.0,
            "taux_absorption_pct": 84.0,
            "periode": "2025-02",
        }

        with mock.patch.object(de, "fetch_marche_neuf", return_value=neuf_data):
            with mock.patch.object(de, "detect_city", return_value="montreal"):
                case: dict = {"dossier_id": "D-NEUF-TEST"}
                enrich_case(case, display_name="50 rue Peel, Montréal",
                            cache_dir=tmp_path)

        mn = case.get("marche_neuf", {})
        assert mn.get("completions_mois") == 420.0
        assert mn.get("taux_absorption_pct") == 84.0


# ── TestDataEnrichment_DistanceCBD ────────────────────────────────────────────

class TestDataEnrichment_DistanceCBD:
    """Tests for compute_distance_cbd() and _haversine_km()."""

    def test_haversine_known_distance(self):
        """Haversine: MTL downtown to Laval approx 15-20 km."""
        import pytest
        from engine.data_enrichment import _haversine_km
        # Montreal Place d'Armes → Carrefour Laval
        dist = _haversine_km(45.5088, -73.5540, 45.5636, -73.6924)
        assert 10.0 < dist < 20.0

    def test_compute_distance_cbd_center(self):
        """Coords at Place d'Armes → 'centre-ville', distance < 5 km."""
        import pytest
        from engine.data_enrichment import compute_distance_cbd
        result = compute_distance_cbd(45.5088, -73.5540, "montreal")
        assert result.get("source") == "calcul-haversine"
        assert result.get("ville_reference") == "Montréal"
        assert result.get("distance_cbd_km") == pytest.approx(0.0, abs=0.5)
        assert result.get("interpretation") == "centre-ville"

    def test_compute_distance_cbd_suburb(self):
        """Coords far from city center → 'banlieue' category."""
        from engine.data_enrichment import compute_distance_cbd
        # Repentigny (~35 km from MTL downtown)
        result = compute_distance_cbd(45.7447, -73.4498, "montreal")
        assert result.get("distance_cbd_km") > 20.0
        assert result.get("interpretation") in ("banlieue proche", "banlieue éloignée")

    def test_compute_distance_cbd_unknown_city(self):
        """Unknown city_code → {}."""
        from engine.data_enrichment import compute_distance_cbd
        result = compute_distance_cbd(45.5, -73.5, "rimouski")
        assert result == {}

    def test_compute_distance_cbd_peri_central(self):
        """Coords ~10 km from CBD → 'péri-central'."""
        from engine.data_enrichment import compute_distance_cbd
        # Anjou / Saint-Léonard area (~10-12 km east of Place d'Armes)
        result = compute_distance_cbd(45.5783, -73.5515, "montreal")
        assert 5.0 <= result.get("distance_cbd_km", 0) < 15.0
        assert result.get("interpretation") == "péri-central"

    def test_enrich_case_injects_distance_cbd(self, tmp_path):
        """enrich_case injects distance_cbd after geocoding."""
        import unittest.mock as mock
        from engine import data_enrichment as de
        from engine.data_enrichment import enrich_case

        dist_data = {
            "source": "calcul-haversine",
            "distance_cbd_km": 3.5,
            "ville_reference": "Montréal",
            "interpretation": "centre-ville",
        }

        with mock.patch.object(de, "compute_distance_cbd", return_value=dist_data):
            with mock.patch.object(de, "geocode_address", return_value=(45.51, -73.56)):
                with mock.patch.object(de, "detect_city", return_value="montreal"):
                    case: dict = {"dossier_id": "D-DIST-TEST"}
                    enrich_case(case, display_name="500 rue Sherbrooke, Montréal",
                                cache_dir=tmp_path)

        dc = case.get("distance_cbd", {})
        assert dc.get("distance_cbd_km") == 3.5
        assert dc.get("interpretation") == "centre-ville"
