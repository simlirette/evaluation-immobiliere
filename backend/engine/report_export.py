"""
Export du rapport d'évaluation : génération .docx et HTML imprimable.
Aucune dépendance sur runtime.py — module isolé.
"""
from __future__ import annotations

import io
import re


HTML_PRINT_TEMPLATE = """\
<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,400;0,500;0,600;1,400&display=swap" rel="stylesheet">
<style>
  *, *::before, *::after {{ box-sizing: border-box; }}
  body {{
    font-family: 'Cormorant Garamond', Georgia, serif;
    font-size: 12pt;
    line-height: 1.65;
    color: #1a1916;
    max-width: 780px;
    margin: 0 auto;
    padding: 2.5cm;
  }}
  .watermark {{
    background: #fff3cd;
    border: 1.5px solid #ffc107;
    border-radius: 4px;
    padding: 8px 14px;
    margin-bottom: 24px;
    font-size: 10pt;
    font-weight: 600;
    color: #856404;
  }}
  h1 {{ font-size: 18pt; font-weight: 600; margin: 0 0 8px; }}
  h2 {{ font-size: 14pt; font-weight: 600; margin: 28px 0 8px;
        border-bottom: 1px solid #e5e2dc; padding-bottom: 4px; }}
  h3 {{ font-size: 12pt; font-weight: 600; margin: 18px 0 6px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 10pt; margin: 12px 0; }}
  th {{ text-align: left; padding: 6px 10px; background: #f5f3ef;
        font-weight: 600; border: 1px solid #ddd9d2; }}
  td {{ padding: 5px 10px; border: 1px solid #ddd9d2; vertical-align: top; }}
  blockquote {{ margin: 12px 0; padding: 8px 14px; border-left: 3px solid #ffc107;
                background: #fffbf0; font-size: 10pt; color: #856404; }}
  ul, ol {{ padding-left: 22px; margin: 8px 0; }}
  li {{ margin: 3px 0; }}
  @media print {{
    @page {{ size: A4; margin: 2.5cm; }}
    body {{ padding: 0; max-width: none; }}
    table {{ page-break-inside: avoid; }}
  }}
</style>
</head>
<body>
<div class="watermark">&#9888; BROUILLON NON CERTIFIÉ — Ce document doit être révisé et signé
par un évaluateur agréé (É.A.) avant toute diffusion.
Supprimer cet avertissement avant certification.</div>
{body}
</body>
</html>
"""


def _generate_html(md_text: str, dossier_id: str) -> str:
    """Convertit le markdown en HTML imprimable avec CSS A4 et watermark."""
    import markdown as md_lib  # type: ignore
    body = md_lib.markdown(md_text, extensions=["tables", "fenced_code"])
    return HTML_PRINT_TEMPLATE.format(
        title=f"Rapport d'évaluation — {dossier_id}",
        body=body,
    )


def _parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    """Retourne liste de (texte, bold, italic) depuis markdown inline."""
    parts: list[tuple[str, bool, bool]] = []
    pattern = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)")
    for m in pattern.finditer(text):
        if m.group(1):
            parts.append((m.group(1), True, True))
        elif m.group(2):
            parts.append((m.group(2), True, False))
        elif m.group(3):
            parts.append((m.group(3), False, True))
        elif m.group(4):
            parts.append((m.group(4), False, False))
    return parts


def _add_inline_paragraph(doc: object, text: str, style: str = "Normal") -> object:  # type: ignore
    """Ajoute un paragraphe avec runs gras/italique inline."""
    p = doc.add_paragraph(style=style)  # type: ignore[attr-defined]
    for chunk, bold, italic in _parse_inline(text):
        run = p.add_run(chunk)
        run.bold = bold
        run.italic = italic
    return p


def _generate_pdf(md_text: str, dossier_id: str) -> bytes:
    """Convertit le markdown en PDF via PyMuPDF fitz.Story avec en-tête et pagination."""
    import fitz  # type: ignore
    import markdown as md_lib  # type: ignore
    import datetime as _dt

    body = md_lib.markdown(md_text, extensions=["tables", "fenced_code"])
    html = f"""<!DOCTYPE html>
<html lang="fr"><head><meta charset="utf-8"><style>
body {{ font-family: Times New Roman, Times, serif; font-size: 11pt; line-height: 1.65; color: #1a1916; margin: 0; padding: 0; }}
.watermark {{ background: #fff8e1; border: 1px solid #b8860b; border-radius: 3px;
    padding: 7px 12px; margin-bottom: 18px; font-size: 9pt; font-weight: bold; color: #7a5500; }}
h1 {{ font-size: 17pt; font-weight: bold; margin: 0 0 10px; }}
h2 {{ font-size: 13pt; font-weight: bold; margin: 22px 0 8px;
    border-bottom: 1px solid #ccc; padding-bottom: 3px; }}
h3 {{ font-size: 11pt; font-weight: bold; margin: 14px 0 5px; }}
table {{ width: 100%; border-collapse: collapse; font-size: 9pt; margin: 10px 0; }}
th {{ text-align: left; padding: 5px 8px; background: #f0ede8;
    font-weight: bold; border: 1px solid #ccc; }}
td {{ padding: 4px 8px; border: 1px solid #ccc; vertical-align: top; }}
ul, ol {{ padding-left: 20px; margin: 6px 0; }}
li {{ margin: 2px 0; }}
blockquote {{ margin: 10px 0; padding: 6px 12px; border-left: 3px solid #b8860b;
    background: #fffbf0; font-size: 9pt; color: #7a5500; }}
</style></head>
<body>
<div class="watermark">&#9888; BROUILLON NON CERTIFI\u00c9 \u2014 Ce document doit \u00eatre r\u00e9vis\u00e9 et sign\u00e9
par un \u00e9valuateur agr\u00e9\u00e9 (\u00c9.A.) avant toute diffusion.</div>
{body}
</body></html>"""

    # Dimensions A4 en points (1 pt = 1/72 po)
    mediabox = fitz.paper_rect("a4")          # 595 × 842 pt
    margin_h = 72                              # ~2.54 cm horizontal
    header_h = 40                              # espace en-tête
    footer_h = 32                              # espace pied
    where = mediabox + (margin_h, header_h, -margin_h, -footer_h)

    # ── Passe 1 : génération du corps ────────────────────────────────────────
    buf = io.BytesIO()
    story = fitz.Story(html=html)
    writer = fitz.DocumentWriter(buf)
    more = True
    while more:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()
    writer.close()

    # ── Passe 2 : en-tête et pied sur chaque page ───────────────────────────
    doc = fitz.open("pdf", buf.getvalue())
    total_pages = doc.page_count
    date_str = _dt.date.today().strftime("%Y-%m-%d")
    grey = (0.55, 0.55, 0.55)
    separator_color = (0.75, 0.75, 0.75)

    for i, page in enumerate(doc):
        w = page.rect.width
        # Ligne séparatrice sous l'en-tête
        page.draw_line(
            fitz.Point(margin_h, header_h - 6),
            fitz.Point(w - margin_h, header_h - 6),
            width=0.5, color=separator_color,
        )
        # Texte en-tête gauche : dossier_id
        page.insert_text(
            fitz.Point(margin_h, header_h - 14),
            dossier_id,
            fontsize=8, color=grey,
        )
        # Texte en-tête droite : date + statut
        label_right = f"{date_str} \u2014 BROUILLON NON CERTIFI\u00c9"
        page.insert_text(
            fitz.Point(w - margin_h - len(label_right) * 4.5, header_h - 14),
            label_right,
            fontsize=8, color=(0.65, 0.35, 0.0),
        )
        # Ligne séparatrice au-dessus du pied
        footer_y = mediabox.y1 - footer_h + 8
        page.draw_line(
            fitz.Point(margin_h, footer_y),
            fitz.Point(w - margin_h, footer_y),
            width=0.5, color=separator_color,
        )
        # Numéro de page centré
        page_label = f"Page {i + 1} / {total_pages}"
        page.insert_text(
            fitz.Point(w / 2 - len(page_label) * 2.5, footer_y + 14),
            page_label,
            fontsize=8, color=grey,
        )

    out = io.BytesIO()
    doc.save(out, garbage=4, deflate=True)
    doc.close()
    return out.getvalue()


def _generate_docx(md_text: str, dossier_id: str) -> bytes:
    """Convertit le markdown en fichier .docx (python-docx)."""
    from docx import Document  # type: ignore
    from docx.shared import Cm, RGBColor  # type: ignore

    doc = Document()

    # Marges 2.5cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Watermark — paragraphe rouge gras en tête
    wm = doc.add_paragraph()
    run = wm.add_run(
        "\u26a0 BROUILLON NON CERTIFIÉ — Ce document doit être révisé et signé par un "
        "évaluateur agréé (É.A.) avant toute diffusion. "
        "Supprimer ce paragraphe avant certification."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    doc.add_paragraph()  # espace après watermark

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading 1 (# mais pas ##)
        if re.match(r"^# [^#]", line):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2 (## mais pas ###)
        if re.match(r"^## [^#]", line):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if re.match(r"^### ", line):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            _add_inline_paragraph(doc, line[2:].strip(), style="List Bullet")
            i += 1
            continue

        # Table (lignes commençant par |)
        if line.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Exclure la ligne de séparation |---|---|
            data_rows = [l for l in table_lines if not re.match(r"^\|[\s\-:|]+\|", l)]
            if not data_rows:
                continue
            try:
                parsed = [[c.strip() for c in row.strip("|").split("|")] for row in data_rows]
                max_cols = max(len(r) for r in parsed)
                table = doc.add_table(rows=len(parsed), cols=max_cols)
                table.style = "Table Grid"
                for row_idx, row_data in enumerate(parsed):
                    for col_idx in range(max_cols):
                        cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text
                        if row_idx == 0:
                            for r in cell.paragraphs[0].runs:
                                r.bold = True
            except Exception:
                # Fallback plain text si table malformée
                for l in data_rows:
                    doc.add_paragraph(l.strip("|").replace("|", "  "), style="Normal")
            doc.add_paragraph()
            continue

        # Ligne vide
        if not line.strip():
            i += 1
            continue

        # Paragraphe normal avec inline bold/italic
        _add_inline_paragraph(doc, line, style="Normal")
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Exports certifiés (T3.5) ──────────────────────────────────────────────────

HTML_CERTIFIED_TEMPLATE = """\
<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,400;0,500;0,600;1,400&display=swap" rel="stylesheet">
<style>
  *, *::before, *::after {{ box-sizing: border-box; }}
  body {{
    font-family: 'Cormorant Garamond', Georgia, serif;
    font-size: 12pt;
    line-height: 1.65;
    color: #1a1916;
    max-width: 780px;
    margin: 0 auto;
    padding: 2.5cm;
  }}
  .cert-header {{
    background: #f0f7f0;
    border: 1.5px solid #4a6b54;
    border-radius: 4px;
    padding: 8px 14px;
    margin-bottom: 24px;
    font-size: 10pt;
    font-weight: 600;
    color: #2d4a35;
  }}
  .signature-bloc {{
    margin-top: 32px;
    padding: 16px;
    border: 1px solid #4a6b54;
    border-radius: 4px;
    background: #f8faf8;
  }}
  .signature-bloc h3 {{ margin: 0 0 12px; font-size: 12pt; color: #2d4a35; }}
  .signature-line {{ margin: 8px 0; font-size: 11pt; }}
  .signature-blank {{ border-bottom: 1px solid #333; height: 40px; margin: 12px 0 4px; }}
  h1 {{ font-size: 18pt; font-weight: 600; margin: 0 0 8px; }}
  h2 {{ font-size: 14pt; font-weight: 600; margin: 28px 0 8px;
        border-bottom: 1px solid #e5e2dc; padding-bottom: 4px; }}
  h3 {{ font-size: 12pt; font-weight: 600; margin: 18px 0 6px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 10pt; margin: 12px 0; }}
  th {{ text-align: left; padding: 6px 10px; background: #f5f3ef;
        font-weight: 600; border: 1px solid #ddd9d2; }}
  td {{ padding: 5px 10px; border: 1px solid #ddd9d2; vertical-align: top; }}
  blockquote {{ margin: 12px 0; padding: 8px 14px; border-left: 3px solid #4a6b54;
                background: #f0f7f0; font-size: 10pt; color: #2d4a35; }}
  ul, ol {{ padding-left: 22px; margin: 8px 0; }}
  li {{ margin: 3px 0; }}
  @media print {{
    @page {{ size: A4; margin: 2.5cm; }}
    body {{ padding: 0; max-width: none; }}
    table {{ page-break-inside: avoid; }}
  }}
</style>
</head>
<body>
<div class="cert-header">✓ RAPPORT D'ÉVALUATION CERTIFIÉ — Signé par {nom_ea} | Permis OEAQ {no_permis} | {date_sig}</div>
{body}
{signature_bloc}
</body>
</html>
"""


def _signature_bloc_html(signature: dict) -> str:
    nom = signature.get("nom_ea", "—")
    no_permis = signature.get("no_permis_oeaq", "—")
    date_sig = signature.get("date_signature", "—")
    return f"""
<div class="signature-bloc">
  <h3>Attestation et signature</h3>
  <div class="signature-line"><strong>Évaluateur agréé :</strong> {nom}</div>
  <div class="signature-line"><strong>N° de permis OEAQ :</strong> {no_permis}</div>
  <div class="signature-line"><strong>Date de signature :</strong> {date_sig}</div>
  <div class="signature-blank"></div>
  <div class="signature-line" style="font-size:9pt;color:#888">_____________________________________________<br>
  Signature de l'évaluateur agréé | Sceau professionnel</div>
</div>"""


def _remove_brouillon_watermark(md_text: str) -> str:
    """Retire les marqueurs BROUILLON du markdown avant génération certifiée."""
    lines = md_text.splitlines()
    cleaned = []
    for line in lines:
        lower = line.lower()
        if ("brouillon non certifié" in lower or
                "brouillon non certifie" in lower or
                "mode dégradé" in lower or
                "mode degrade" in lower or
                "produit par assistant ia" in lower):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def generate_certified_html(md_text: str, dossier_id: str, signature: dict) -> str:
    """Génère le HTML certifié — sans filigrane BROUILLON, avec bloc signature.

    Action explicite É.A. uniquement — jamais appelée automatiquement.
    """
    import markdown as md_lib  # type: ignore
    clean_md = _remove_brouillon_watermark(md_text)
    body = md_lib.markdown(clean_md, extensions=["tables", "fenced_code"])
    nom = signature.get("nom_ea", "—")
    no_permis = signature.get("no_permis_oeaq", "—")
    date_sig = signature.get("date_signature", "—")
    return HTML_CERTIFIED_TEMPLATE.format(
        title=f"Rapport certifié — {dossier_id}",
        nom_ea=nom,
        no_permis=no_permis,
        date_sig=date_sig,
        body=body,
        signature_bloc=_signature_bloc_html(signature),
    )


def generate_certified_pdf(md_text: str, dossier_id: str, signature: dict) -> bytes:
    """Génère le PDF certifié — sans filigrane, avec en-tête certifié et bloc signature.

    Action explicite É.A. uniquement.
    """
    import fitz  # type: ignore
    import markdown as md_lib  # type: ignore
    import datetime as _dt

    clean_md = _remove_brouillon_watermark(md_text)
    nom = signature.get("nom_ea", "—")
    no_permis = signature.get("no_permis_oeaq", "—")
    date_sig = signature.get("date_signature", "—")

    sig_html = f"""
<div style="margin-top:32px;padding:14px;border:1px solid #4a6b54;border-radius:4px;background:#f8faf8;">
<strong>Attestation et signature</strong><br/>
Évaluateur agréé : {nom}<br/>
N° de permis OEAQ : {no_permis}<br/>
Date de signature : {date_sig}<br/>
<div style="border-bottom:1px solid #333;height:40px;margin:12px 0 4px;"></div>
<span style="font-size:9pt;color:#888">Signature É.A. | Sceau professionnel</span>
</div>"""

    body = md_lib.markdown(clean_md, extensions=["tables", "fenced_code"])
    html = f"""<!DOCTYPE html>
<html lang="fr"><head><meta charset="utf-8"><style>
body {{ font-family: Times New Roman, Times, serif; font-size: 11pt; line-height: 1.65; color: #1a1916; margin: 0; padding: 0; }}
.cert-header {{ background: #f0f7f0; border: 1px solid #4a6b54; border-radius: 3px;
    padding: 7px 12px; margin-bottom: 18px; font-size: 9pt; font-weight: bold; color: #2d4a35; }}
h1 {{ font-size: 17pt; font-weight: bold; margin: 0 0 10px; }}
h2 {{ font-size: 13pt; font-weight: bold; margin: 22px 0 8px;
    border-bottom: 1px solid #ccc; padding-bottom: 3px; }}
h3 {{ font-size: 11pt; font-weight: bold; margin: 14px 0 5px; }}
table {{ width: 100%; border-collapse: collapse; font-size: 9pt; margin: 10px 0; }}
th {{ text-align: left; padding: 5px 8px; background: #f0ede8;
    font-weight: bold; border: 1px solid #ccc; }}
td {{ padding: 4px 8px; border: 1px solid #ccc; vertical-align: top; }}
ul, ol {{ padding-left: 20px; margin: 6px 0; }}
li {{ margin: 2px 0; }}
blockquote {{ margin: 10px 0; padding: 6px 12px; border-left: 3px solid #4a6b54;
    background: #f0f7f0; font-size: 9pt; color: #2d4a35; }}
</style></head>
<body>
<div class="cert-header">✓ RAPPORT CERTIFIÉ — {nom} | Permis {no_permis} | {date_sig}</div>
{body}
{sig_html}
</body></html>"""

    mediabox = fitz.paper_rect("a4")
    margin_h = 72
    header_h = 40
    footer_h = 32
    where = mediabox + (margin_h, header_h, -margin_h, -footer_h)

    buf = io.BytesIO()
    story = fitz.Story(html=html)
    writer = fitz.DocumentWriter(buf)
    more = True
    while more:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()
    writer.close()

    doc = fitz.open("pdf", buf.getvalue())
    total_pages = doc.page_count
    grey = (0.4, 0.5, 0.4)
    separator_color = (0.6, 0.75, 0.6)

    for i, page in enumerate(doc):
        w = page.rect.width
        page.draw_line(
            fitz.Point(margin_h, header_h - 6),
            fitz.Point(w - margin_h, header_h - 6),
            width=0.5, color=separator_color,
        )
        page.insert_text(
            fitz.Point(margin_h, header_h - 14),
            dossier_id,
            fontsize=8, color=grey,
        )
        cert_label = f"{date_sig} — CERTIFIÉ É.A. {nom[:25]}"
        page.insert_text(
            fitz.Point(w - margin_h - len(cert_label) * 4.5, header_h - 14),
            cert_label,
            fontsize=8, color=(0.18, 0.42, 0.18),
        )
        footer_y = mediabox.y1 - footer_h + 8
        page.draw_line(
            fitz.Point(margin_h, footer_y),
            fitz.Point(w - margin_h, footer_y),
            width=0.5, color=separator_color,
        )
        page_label = f"Page {i + 1} / {total_pages}"
        page.insert_text(
            fitz.Point(w / 2 - len(page_label) * 2.5, footer_y + 14),
            page_label,
            fontsize=8, color=grey,
        )

    out = io.BytesIO()
    doc.save(out, garbage=4, deflate=True)
    doc.close()
    return out.getvalue()


def generate_certified_docx(md_text: str, dossier_id: str, signature: dict) -> bytes:
    """Génère le .docx certifié — sans filigrane BROUILLON, avec bloc signature.

    Action explicite É.A. uniquement.
    """
    from docx import Document  # type: ignore
    from docx.shared import Cm, RGBColor  # type: ignore

    nom = signature.get("nom_ea", "—")
    no_permis = signature.get("no_permis_oeaq", "—")
    date_sig = signature.get("date_signature", "—")

    clean_md = _remove_brouillon_watermark(md_text)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # En-tête certifié (vert)
    cert_p = doc.add_paragraph()
    cert_run = cert_p.add_run(
        f"✓ RAPPORT D'ÉVALUATION CERTIFIÉ — {nom} | Permis OEAQ {no_permis} | {date_sig}"
    )
    cert_run.bold = True
    cert_run.font.color.rgb = RGBColor(0x2D, 0x4A, 0x35)
    doc.add_paragraph()

    lines = clean_md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if re.match(r"^# [^#]", line):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^## ", line):
            doc.add_heading(line[3:].strip(), level=2)
        elif re.match(r"^### ", line):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip():
            _add_inline_paragraph(doc, line, style="Normal")
        else:
            doc.add_paragraph()
        i += 1

    # Bloc signature
    doc.add_page_break()
    doc.add_heading("Attestation et signature", level=2)
    for label, value in [
        ("Évaluateur agréé", nom),
        ("N° de permis OEAQ", no_permis),
        ("Date de signature", date_sig),
        ("Dossier", dossier_id),
    ]:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label} : ")
        run_label.bold = True
        p.add_run(value)
    doc.add_paragraph()
    sig_line = doc.add_paragraph("_" * 60)
    sig_line.add_run("\nSignature de l'évaluateur agréé | Sceau professionnel")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
